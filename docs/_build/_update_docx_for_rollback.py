"""
Replace Part 11.12 (the v2 CMS section) in the DOCX with a small note about
the CMS being prototyped then rolled back per operator preference.

Strategy:
- Find the heading "11.12 v2 Update..." and delete everything from that heading
  to the end of the document (since 11.12 was appended at the very end).
- Append a new short note describing the rollback.
"""
import sys, os
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
doc = Document(IN_PATH)

body = doc.element.body

# 1. Find the index of the paragraph that contains "11.12 v2 Update"
v2_start_idx = None
for i, p in enumerate(doc.paragraphs):
    if '11.12 v2 Update' in p.text or '11.12 v2 update' in p.text.lower():
        v2_start_idx = i
        break

if v2_start_idx is None:
    print("Could not find Part 11.12 heading — no changes made.")
else:
    print(f"Found Part 11.12 at paragraph index {v2_start_idx}")

    # 2. Collect all elements to delete (the paragraph at v2_start_idx + everything after)
    # until end of body
    paragraphs = doc.paragraphs
    total_paras = len(paragraphs)

    # Get the element of the v2 paragraph
    v2_para_elem = paragraphs[v2_start_idx]._p

    # Walk siblings after v2_para_elem and remove them
    sibling = v2_para_elem
    removed = 0
    while sibling is not None:
        next_sibling = sibling.getnext()
        # Don't remove sectPr (last element is section properties)
        if sibling.tag == qn('w:sectPr'):
            # This is the section properties, stop
            sibling = None
            break
        # Remove the element from its parent
        parent = sibling.getparent()
        if parent is not None:
            parent.remove(sibling)
            removed += 1
        sibling = next_sibling

    print(f"Removed {removed} elements from Part 11.12 to end of body")

# 3. Append a small note about the rollback at end of document
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0B, 0x11, 0x20)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x4B, 0x5E)
    return h

def add_para(text, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.italic = italic
    return p

# Add update note at end
add_heading('11.12 v2 Rollback Note (2026-07-04)', level=2)

add_heading('11.12.1 What Changed', level=3)
for s in [
    'A separate CMS admin form at /super-admin/pricing-settings was prototyped, then REMOVED per operator preference. The operator controls all plan content from /super-admin/plans.',
    'All marketing copy (hero, trust stats, features, addons, FAQs, footer) is now hard-coded directly in resources/views/pricing/index.blade.php.',
    'WhatsApp CTAs remain REMOVED from the public page (retained from the v2 change that the operator asked for).',
    'The only CTA per pricing card is now "احجز ديمو مجاني" pointing to the configured demo URL.',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading('11.12.2 Removed Files', level=3)
for s in [
    'app/Models/PricingSetting.php',
    'app/Http/Controllers/Landlord/PricingSettingsController.php',
    'resources/views/landlord/pricing-settings.blade.php',
    'database/migrations/2026_07_04_000001_create_pricing_settings_table.php',
    'database/seeders/PricingSettingSeeder.php',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading('11.12.3 Reverted Files (now back to hard-coded v1 behavior)', level=3)
for s in [
    'routes/web.php — /super-admin/pricing-settings routes removed',
    'resources/views/landlord/layout.blade.php — sidebar item "إعدادات صفحة الأسعار" removed',
    'app/Http/Controllers/PricingController.php — back to simple Plan::where(...)->get() pattern; settings removed',
    'resources/views/pricing/index.blade.php — hard-coded sections, no $settings; WhatsApp still removed',
    '.env.example + .env.production.example — KAYYAN_WHATSAPP_NUMBER removed',
    'config/pricing.php — whatsapp_number key removed',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading('11.12.4 DB Action Taken', level=3)
add_para('php artisan tinker → Schema::drop("pricing_settings") — table dropped (migration file was also deleted).', italic=True)

add_heading('11.12.5 How To Edit Marketing Content Going Forward', level=3)
add_para('Edit resources/views/pricing/index.blade.php directly. Each section has hard-coded arrays (stats, features, addons, faqs) at the top of its section block. No artisan command needed — Blade recompiles on demand.')
add_para('To edit plan-level data (name, price, features, slugs, is_active), use /super-admin/plans (unchanged).', italic=True)

doc.save(IN_PATH)
print(f'\nSaved: {IN_PATH}')
print(f'Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB')

# Verify
doc2 = Document(IN_PATH)
print('\n=== Part 11 sub-sections ===')
capture = False
for p in doc2.paragraphs:
    if 'Part 11' in p.text:
        capture = True
    elif capture and p.style.name == 'Heading 1':
        break
    elif capture and p.style.name in ('Heading 2', 'Heading 3'):
        prefix = '  ' if p.style.name == 'Heading 2' else '    '
        print(f'{prefix}{p.text}')
