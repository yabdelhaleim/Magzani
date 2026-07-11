"""
Export DOCX → clean TXT for easy reading in any editor.
Run from docs/ directory:
    python _export_docx_to_txt.py
"""
import os, sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from docx import Document
from docx.shared import RGBColor

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
OUT_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.txt"

doc = Document(IN_PATH)

lines = []
lines.append("=" * 80)
lines.append("  Magzani ERP — Comprehensive Technical Documentation")
lines.append("  (Plain-text export of Magzani_ERP_Full_Documentation.docx)")
lines.append(f"  File size: {os.path.getsize(IN_PATH) / 1024:.1f} KB")
lines.append(f"  Sections: {len(doc.sections)}, Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
lines.append("=" * 80)
lines.append("")

# Track which table we're at (for interweaving)
table_iter = iter(doc.tables)
table_idx = 0
all_tables = doc.tables

def render_table(table):
    """Render a docx table as ASCII."""
    out = []
    for row in table.rows:
        cells_text = []
        for cell in row.cells:
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            # Truncate long cells
            if len(text) > 80:
                text = text[:77] + "..."
            cells_text.append(text)
        out.append(" | ".join(cells_text))
    return "\n".join(out)

# Iterate paragraphs in body order; tables are interleaved
body = doc.element.body
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        # Find matching paragraph by index
        for p in doc.paragraphs:
            if p._p is child:
                text = p.text.strip()
                style = p.style.name
                if not text:
                    lines.append("")
                    continue
                if style == 'Heading 1':
                    lines.append("")
                    lines.append("=" * 80)
                    lines.append(f"# {text}")
                    lines.append("=" * 80)
                    lines.append("")
                elif style == 'Heading 2':
                    lines.append("")
                    lines.append(f"## {text}")
                    lines.append("")
                elif style == 'Heading 3':
                    lines.append("")
                    lines.append(f"### {text}")
                    lines.append("")
                else:
                    lines.append(text)
                break
    elif tag == 'tbl':
        # Find matching table by index
        for t in all_tables:
            if t._tbl is child:
                lines.append("")
                lines.append("┌" + "─" * 78 + "┐")
                lines.append(render_table(t))
                lines.append("└" + "─" * 78 + "┘")
                lines.append("")
                break

with open(OUT_PATH, 'w', encoding='utf-8') as f:
    f.write("\n".join(lines))

print(f"Exported to: {OUT_PATH}")
print(f"Size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
print(f"Lines: {len(lines)}")
