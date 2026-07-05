"""
Post-build fix v2:
- Rename "Part 3: Database Schema" → "Part 5: Database Schema" (its 3.1/3.2 sections → 5.1/5.2)
- Shift Part 5..9 forward by 1 → Part 6..10 (with their inner 5.x, 6.x, etc. renumbered)

Strategy:
- Find each "Part N: ..." heading and store its text + original number
- Build a renumber map: N → N+1 for N in {5,6,7,8,9}
- Database Schema N=3 → new N=5 (with 3.x → 5.x)
- Apply renames in reverse order (so we don't double-replace).
"""
from docx import Document
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"

doc = Document(IN_PATH)

# Step 1: Rename duplicate Part 3: Database Schema → Part 5: Database Schema
# Step 2: Shift all Part 5..9 → 6..10 (which we already did in part v1, but verify)
# Step 3: Also rename inner 3.1, 3.2, 3.3 → 5.1, 5.2, 5.3 (DB Schema section numbers)

# Process top-level H1 renames
def rename_paragraph_text(p, new_text):
    """Replace all run text with new_text while preserving the first run's formatting."""
    if not p.runs:
        run = p.add_run(new_text)
        return
    first = True
    for run in p.runs:
        if first:
            run.text = new_text
            first = False
        else:
            run.text = ''

# Step 1: Find all H1 paragraphs and their texts
h1_indices = []  # list of (idx, paragraph, text)
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1':
        h1_indices.append((i, p))

# Step 2: Apply title renames
title_renames = {
    'Part 3: Database Schema':               'Part 5: Database Schema',
    'Part 6: Complete Routes Reference':     'Part 7: Complete Routes Reference',
    'Part 7: Business Logic Deep-Dive':      'Part 8: Business Logic Deep-Dive',
    'Part 8: Recent Changes & Known Issues': 'Part 9: Recent Changes & Known Issues',
    'Part 9: Quick Reference Card':          'Part 10: Quick Reference Card',
    # Also shift Modules / Business Logic if needed:
    'Part 5: Modules (Deep Dive)':           'Part 6: Modules (Deep Dive)',
    # But these might already be Part 6 in the actual file. Let's check first.
}

for _, p in h1_indices:
    txt = p.text.strip()
    if txt in title_renames:
        rename_paragraph_text(p, title_renames[txt])
        print(f"  H1: {txt} → {title_renames[txt]}")

# Step 3: Apply H2/H3 section renumber inside DB Schema
# The DB Schema sections are 3.1 → 3.14 currently. They'll become 5.1 → 5.14.
# We need to walk paragraphs in order and rename only those that fall between
# "Part 5: Database Schema" (or "Part 3: Database Schema" before rename) and the next H1.
all_paras = doc.paragraphs

# Find the indices of H1 boundaries
h1_positions = [i for i, p in enumerate(all_paras) if p.style.name == 'Heading 1']

# Determine boundaries for DB schema
db_start = None
db_end = None
for idx_pos, idx in enumerate(h1_positions):
    p = all_paras[idx]
    if 'Database Schema' in p.text:
        db_start = idx
        if idx_pos + 1 < len(h1_positions):
            db_end = h1_positions[idx_pos + 1]
        else:
            db_end = len(all_paras)
        break

if db_start is not None:
    print(f"\n  Database Schema spans paragraphs {db_start}..{db_end}")
    # Iterate paragraphs in DB schema range
    section_prefix_replacements = [
        # old prefix → new prefix
        ('3.1 Central',         '5.1 Central'),
        ('3.2 Tenant Database', '5.2 Tenant Database'),
        ('3.2.1 Auth',          '5.2.1 Auth'),
        ('3.2.2 Catalog',       '5.2.2 Catalog'),
        ('3.2.3 Warehouses',    '5.2.3 Warehouses'),
        ('3.2.4 Sales',         '5.2.4 Sales'),
        ('3.2.5 Purchases',     '5.2.5 Purchases'),
        ('3.2.6 Partners',      '5.2.6 Partners'),
        ('3.2.7 Payments',      '5.2.7 Payments'),
        ('3.2.8 Stock',         '5.2.8 Stock'),
        ('3.2.9 Inventory',     '5.2.9 Inventory'),
        ('3.2.10 Cash',         '5.2.10 Cash'),
        ('3.2.11 Manufacturing','5.2.11 Manufacturing'),
        ('3.2.12 Wood',         '5.2.12 Wood'),
        ('3.2.13 Accounting',   '5.2.13 Accounting'),
        ('3.2.14 Misc',         '5.2.14 Misc'),
        ('3.3 Key Model',       '5.3 Key Model'),
    ]
    for i in range(db_start, db_end):
        p = all_paras[i]
        if p.style.name in ('Heading 2', 'Heading 3'):
            for old, new in section_prefix_replacements:
                if p.text.strip().startswith(old):
                    # Replace the whole text with new
                    rename_paragraph_text(p, p.text.strip().replace(old, new, 1))
                    break

doc.save(IN_PATH)
print(f"\nSaved: {IN_PATH}")
print(f"Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB")

# Verify
doc2 = Document(IN_PATH)
print("\n=== Final H1 list ===")
for p in doc2.paragraphs:
    if p.style.name == 'Heading 1':
        print(f'  • {p.text}')
