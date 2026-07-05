"""
Magzani ERP Doc - Part 2
Tech Stack + File Tree + Database Schema Overview
"""
import sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from _build_doc_part1 import (
    doc, add_heading, add_para, add_para_rtl, add_table, add_code_block,
    add_hr, add_page_break, OUT_PATH,
)
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# PART 2 - TECH STACK
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 2: Tech Stack', level=1)

add_heading(doc, '2.1 Backend Stack', level=2)
add_table(doc,
    ['Component', 'Version', 'Purpose'],
    [
        ['PHP', '^8.2', 'Runtime'],
        ['Laravel Framework', '^12.0', 'Core MVC framework'],
        ['Stancl Tenancy', '^3.10', 'Multi-tenant database isolation'],
        ['Livewire', '^3.5', 'Reactive full-stack components (POS)'],
        ['Laravel Sanctum', '^4.0', 'API token authentication'],
        ['Laravel Reverb', '^1.0', 'WebSocket broadcasting'],
        ['Laravel Tinker', '^2.10', 'REPL'],
        ['GuzzleHTTP', '^7.2', 'HTTP client'],
        ['Doctrine DBAL', '^3.10', 'Database abstraction (used by backup)'],
        ['Predis', '^3.3', 'Redis client'],
        ['Barryvdh DomPDF', '^3.1', 'PDF generation (invoices, statements)'],
        ['Maatwebsite Excel', '^3.1', 'Excel import/export'],
        ['Spatie Laravel-Backup', '^9.0|^10.0', 'Database backups'],
    ],
    col_widths=[Cm(4.5), Cm(3), Cm(8)],
)

add_heading(doc, '2.2 Frontend Stack', level=2)
add_table(doc,
    ['Component', 'Version', 'Purpose'],
    [
        ['Vue', '^3.5.33', 'Declared in package.json (no compiled components currently)'],
        ['Laravel Echo', '^1.16.1', 'WebSocket client (for Reverb)'],
        ['Pusher JS', '^8.4.0', 'Pusher-compatible realtime client'],
        ['Vite', '^5.0.0', 'Asset bundler'],
        ['Laravel Vite Plugin', '^1.0.0', 'Vite integration'],
        ['Tailwind CSS', 'CDN', 'Utility classes via CDN'],
        ['Bootstrap 5 RTL', 'CDN', 'RTL CSS framework via CDN'],
        ['Alpine.js', 'CDN', 'Reactivity helpers'],
        ['Font Awesome', 'CDN', 'Icon set'],
        ['Rubik / Tajawal / Cairo', 'Google Fonts', 'Arabic + Latin typography'],
    ],
    col_widths=[Cm(4.5), Cm(3), Cm(8)],
)

add_heading(doc, '2.3 Dev Dependencies', level=2)
add_table(doc,
    ['Component', 'Purpose'],
    [
        ['PHPUnit ^11.5', 'Unit/feature testing'],
        ['Laravel Pint ^1.0', 'Code style fixer (PSR-12)'],
        ['Mockery ^1.4.4', 'Mocking library'],
        ['FakerPHP ^1.9.1', 'Fake data generator'],
        ['Spatie Ignition ^2.4', 'Error page UI'],
        ['Nuno Maduro Collision ^8.5', 'Pretty test runner output'],
        ['Laravel Sail ^1.18', 'Docker dev environment'],
        ['Pest Plugin (allow)', 'Permitted plugin in composer'],
    ],
    col_widths=[Cm(5.5), Cm(10)],
)

add_heading(doc, '2.4 Database Configuration', level=2)
add_para(doc, 'Default connection is configured in .env via DB_CONNECTION. Production-style config file (.env.production.example) is provided.')
add_table(doc,
    ['Setting', 'Default / Production Note'],
    [
        ['DB_CONNECTION', 'env-driven, default mysql'],
        ['Central DB name', '"laravel" by default in .env.example'],
        ['Tenant DB naming', 'tenant{tenant_id} (prefix tenant + UUID, no suffix)'],
        ['Migrations table', 'migrations (standard)'],
        ['Redis client', 'phpredis (multiple connections configured)'],
        ['Cache', 'file driver (config/cache.php)'],
        ['Queue', 'sync (config/queue.php) — can switch to redis/database'],
        ['Session', 'file driver'],
        ['Filesystem disk', 'local (suffixed per tenant)'],
    ],
    col_widths=[Cm(5), Cm(11)],
)

# ============================================================
# PART 3 - FILE TREE
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 3: Complete File Tree', level=1)

add_heading(doc, '3.1 Root Directory', level=2)
add_code_block(doc, """C:\\MAGZANIV6\\Magzani\\
├── app/                  # Application source
├── bootstrap/            # Laravel bootstrap (app, cache)
├── config/               # 15 config files incl. tenancy, auth, livewire
├── database/             # Migrations, seeders, factories
├── docs/                 # THIS DOCUMENT lives here
├── lang/                 # i18n (ar/modules.php)
├── public/               # index.php + assets
├── resources/            # Views + raw assets
├── routes/               # web.php, api.php, tenant.php, etc.
├── storage/              # logs, cache, framework
├── tests/                # PHPUnit tests
├── vendor/               # Composer packages
├── .env / .env.example   # Env config
├── artisan               # CLI entry
├── composer.json         # PHP deps
├── package.json          # JS deps
├── vite.config.js        # Asset bundler config
└── Multiple *.md docs    # Existing Arabic / English docs""")

add_heading(doc, '3.2 app/ Directory Structure', level=2)
add_code_block(doc, """app/
├── Console/
│   └── Kernel.php
├── Enums/                # 5 PHP enums
│   ├── AccountTypeEnum.php
│   ├── JournalEntrySource.php
│   ├── JournalEntryStatus.php
│   ├── NormalBalance.php
│   └── PaymentTerms.php
├── Events/               # Stancl + custom events
├── Exceptions/           # Handler
├── Exports/              # 10 Maatwebsite export classes
├── Http/
│   ├── Controllers/
│   │   ├── Accounting/   # 13 controllers
│   │   ├── Landlord/     # SuperAdminController
│   │   ├── Auth/         # Login + Register
│   │   └── (40+ top-level controllers)
│   └── Middleware/       # Custom middleware (feature gate, etc.)
├── Jobs/                 # Queued jobs
├── Listeners/            # Event listeners
├── Livewire/             # PosPanel + PosReturnPanel
├── Models/               # 74 Eloquent models
├── Notifications/       # Laravel notifications
├── Observers/            # TenantObserver, etc.
├── Policies/             # ProductPolicy, SettingsPolicy
├── Providers/            # 7 service providers
├── Services/             # 31 top-level + 13 Accounting services
└── Traits/               # Reusable traits""")

add_heading(doc, '3.3 database/ Directory Structure', level=2)
add_code_block(doc, """database/
├── factories/            # 4 factories: User, Product, Warehouse, MO
├── migrations/           # LANDLORD migrations
│   ├── 2019_09_15_000010_create_tenants_table.php
│   ├── 2019_09_15_000020_create_domains_table.php
│   ├── 2026_06_01_000001_create_plans_table.php
│   ├── 2026_06_05_000001_create_plan_features_table.php
│   └── 2026_06_05_000002_add_subscription_dates_to_tenants_table.php
├── migrations/tenant/    # 70+ TENANT migrations
│   ├── 2024_01_01_000001_create_users_table.php
│   ├── 000002-000005      # roles + permissions + pivots
│   ├── 000006-...          # warehouses, categories, products
│   ├── .../sales_invoices, sales_returns, ...
│   ├── .../purchase_invoices, purchase_returns, ...
│   ├── .../manufacturing_orders, manufacturing_costs, ...
│   ├── .../accounts, journal_entries, fiscal_* (accounting)
│   ├── .../wood_stocks, wood_dispensings, raw_material_templates
│   └── .../warehouse_inbound_orders, warehouse_outbound_orders, ...
└── seeders/              # 13 seeders
    ├── DatabaseSeeder.php
    ├── PermissionAndRoleSeeder.php       # 22 KB comprehensive RBAC
    ├── DefaultChartOfAccountsSeeder.php  # 12 KB COA
    ├── TestDataSeeder.php                # 21 KB fixtures
    ├── PlanSeeder.php                    # Subscription plans
    └── (ProductSeeder, CustomerSeeder, etc.)""")

add_heading(doc, '3.4 resources/views/ Directory', level=2)
add_code_block(doc, """resources/views/
├── layouts/
│   ├── app.blade.php        # 71KB - main tenant RTL shell
│   └── theme-css.blade.php  # 19KB - design tokens
├── auth/                    # login, register
├── Dashboard/
│   └── dashboard.blade.php
├── landlord/                # Glass-morphism dark theme
│   ├── layout.blade.php
│   ├── dashboard.blade.php
│   ├── plans/{index,create,edit}.blade.php
│   └── tenants/{index,create,edit}.blade.php
├── livewire/                # POS panels
│   ├── pos-panel.blade.php        # 50KB
│   └── pos-return-panel.blade.php # 11KB
├── customers/, suppliers/         # CRUD + statements
├── products/                      # + barcode, pricing, units
├── categories/
├── invoices/
│   ├── sales/{create,edit,index,show,print,store,thermal_receipt}
│   ├── purchases/{create,edit,index,show,print}
│   ├── sales-returns/
│   └── purchase-returns/
├── pos/{history,settings,shift-close,shift-open,x-report,z-report}
├── stock-counts/
├── transfers/, warehouses/        # + partials, warehouse-history
├── manufacturing/                 # + wood-stocks, wood-dispensings
├── manufacturing-orders/          # + raw-materials
├── Movements/                     # inventory-movements index
├── permissions/                   # RBAC management
├── settings/                      # admin only
├── reports/                       # + financial, inventory, profit-loss, wood-*
├── users/                         # admin user mgmt
├── accounting/                    # coa/, dashboard, expenses, fiscal/,
│                                  # fixed-assets/, journal/, posting-failures/,
│                                  # recurring/, reports/, settings/, setup/,
│                                  # vouchers/, treasury
├── plan/upgrade.blade.php
├── errors/{500,http}.blade.php
├── welcome.blade.php            # 27KB - unused default Laravel
└── welcome_central.blade.php    # 25KB - landlord landing""")

add_heading(doc, '3.5 routes/ Directory', level=2)
add_table(doc,
    ['File', 'Lines', 'Scope'],
    [
        ['routes/web.php', '38', 'Landlord only: /super-admin/*'],
        ['routes/api.php', '20', 'Sanctum minimal API (only /api/user)'],
        ['routes/tenant.php', '619', 'ALL tenant routes (inside middleware group)'],
        ['routes/console.php', '-', 'Laravel default scheduler'],
        ['routes/channels.php', '163 bytes', 'Laravel default broadcasting channels (none custom)'],
    ],
    col_widths=[Cm(5), Cm(2), Cm(9)],
)

add_heading(doc, '3.6 config/ Directory', level=2)
add_table(doc,
    ['Config File', 'Purpose / Notable Settings'],
    [
        ['app.php', 'Timezone UTC, locale en (UI Arabic), 7 custom providers'],
        ['auth.php', 'web session guard, eloquent User provider, 60-min password reset'],
        ['broadcasting.php', 'Standard Laravel config'],
        ['cache.php', 'file driver, tag_base "tenant"'],
        ['cors.php', 'Standard CORS'],
        ['database.php', 'mysql default; redis with phpredis client'],
        ['filesystems.php', 'local + public disks (tenancy suffixes them)'],
        ['hashing.php', 'bcrypt default'],
        ['logging.php', 'stack channel'],
        ['mail.php', 'SMTP defaults (mailpit in dev)'],
        ['queue.php', 'sync default'],
        ['reverb.php', 'Reverb WebSocket server (port 8080)'],
        ['sanctum.php', 'API tokens'],
        ['services.php', '3rd-party credentials placeholders'],
        ['session.php', 'file driver, 120-min lifetime'],
        ['tenancy.php', '234 lines — central tenant config (see Part 2)'],
        ['view.php', 'view paths'],
    ],
    col_widths=[Cm(4), Cm(12)],
)

add_heading(doc, '3.7 Lang / Tests / Storage / Public', level=2)
add_code_block(doc, """lang/
└── ar/modules.php         # 405 bytes - Arabic module name labels

tests/                     # PHPUnit skeleton (probably Feature + Unit)

storage/
├── app/                   # tenant-suffixed
├── app/public/            # tenant-suffixed (Laravel <11) 
├── framework/
│   ├── cache/data/
│   ├── sessions/
│   ├── testing/
│   └── views/
└── logs/laravel.log

public/
├── index.php              # Laravel entry
├── .htaccess              # Apache rewrite rules
├── favicon.ico
├── robots.txt
├── css/                   # compiled assets folder
└── build/                 # vite build output""")

doc.save(OUT_PATH)
print(f"Part 2 appended → {OUT_PATH}")
