"""
Magzani ERP - Comprehensive Technical Documentation Generator
Generates a complete DOCX reference document for the Magzani project.

Usage:
    python _build_doc.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"

# ============================================================
# Helpers
# ============================================================

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0B, 0x11, 0x20)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x4B, 0x5E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_para_rtl(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc, headers, rows, header_bg='1E3A8A', header_fg=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = header_fg
        run.font.name = 'Calibri'
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_borders(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
            set_cell_borders(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================
# Build document
# ============================================================

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ----------------------------------------------------------
# COVER PAGE
# ----------------------------------------------------------
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(140)

r = cover.add_run('Magzani ERP')
r.bold = True
r.font.size = Pt(48)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x0B, 0x11, 0x20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('Comprehensive Technical Documentation')
sr.font.size = Pt(22)
sr.font.name = 'Calibri'
sr.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
sr.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Box with project info
info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Project Name', 'Magzani ERP'),
    ('Codebase Path', r'C:\MAGZANIV6\Magzani'),
    ('Branch', 'main'),
    ('Documentation Date', '2026-07-04'),
    ('Stack', 'Laravel 12 + Livewire 3 + Stancl Tenancy 3.x'),
    ('Type', 'Multi-Tenant SaaS ERP (POS / Inventory / Accounting / Manufacturing)'),
]
for i, (k, v) in enumerate(info_data):
    cells = info_table.rows[i].cells
    cells[0].text = ''
    cells[1].text = ''
    p0 = cells[0].paragraphs[0]
    rk = p0.add_run(k)
    rk.bold = True
    rk.font.size = Pt(11)
    rk.font.name = 'Calibri'
    rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cells[0], '1E3A8A')
    set_cell_borders(cells[0])
    p1 = cells[1].paragraphs[0]
    rv = p1.add_run(v)
    rv.font.size = Pt(11)
    rv.font.name = 'Calibri'
    set_cell_bg(cells[1], 'F0F4FF')
    set_cell_borders(cells[1])
    cells[0].width = Cm(5)
    cells[1].width = Cm(10)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run('This document is the SINGLE SOURCE OF TRUTH for the Magzani codebase.\n'
                  'Always consult this file before exploring the project or answering questions.')
nr.font.size = Pt(11)
nr.italic = True
nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
nr.font.name = 'Calibri'

add_page_break(doc)

# ----------------------------------------------------------
# TABLE OF CONTENTS
# ----------------------------------------------------------
add_heading(doc, 'Table of Contents', level=1)

# Word TOC field
toc_para = doc.add_paragraph()
toc_run = toc_para.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = r'TOC \o "1-3" \h \z \u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
toc_run._r.append(fldChar1)
toc_run._r.append(instrText)
toc_run._r.append(fldChar2)
toc_run._r.append(fldChar3)

toc_hint = doc.add_paragraph()
hr = toc_hint.add_run('Tip: After opening this document in Microsoft Word, right-click the TOC above and choose "Update Field" to populate page numbers.')
hr.italic = True
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

add_page_break(doc)

# ============================================================
# PART 1 — PROJECT OVERVIEW
# ============================================================
add_heading(doc, 'Part 1: Project Overview', level=1)

add_heading(doc, '1.1 What is Magzani?', level=2)
add_para(doc,
    'Magzani ERP is a comprehensive multi-tenant Inventory, Sales, Purchases, POS, Manufacturing and Accounting '
    'system. It is built on Laravel 12 and uses Stancl Tenancy 3.x for database-per-tenant multi-tenancy.')
add_para_rtl(doc,
    'Magzani هو نظام ERP متكامل متعدد المستأجرين (Multi-Tenant) لإدارة المخزون والمبيعات '
    'والمشتريات ونقاط البيع والتصنيع والمحاسبة. مبني على Laravel 12 ويستخدم Stancl Tenancy 3.x '
    'لعزل كل عميل في قاعدة بيانات مستقلة.')

add_heading(doc, '1.2 Business Domain', level=2)
add_para(doc, 'The system targets small-to-medium businesses and supports:')
for b in ['Inventory management across multiple warehouses',
          'Point of Sale with shift-based cash management',
          'Sales & purchase invoicing (multi-tax, multi-discount)',
          'Returns processing (sales & purchase)',
          'Manufacturing / BOM-based production orders',
          'Full accounting: Chart of Accounts, Journal Entries, Fiscal Periods, Fixed Assets',
          'Customer & Supplier ledgers with statements and aging',
          'Multi-warehouse transfers, stock counts, adjustments',
          'Reporting: P&L, Balance Sheet, Trial Balance, Inventory, Wood stock, Manufacturing cost']:
    p = doc.add_paragraph(b, style='List Bullet')

add_heading(doc, '1.3 Target Users', level=2)
add_table(doc,
    ['Role', 'Arabic', 'Description', 'Access Surface'],
    [
        ['Super Admin', 'مدير النظام العام', 'SaaS platform owner', 'routes/web.php (/super-admin/*) — landlord DB'],
        ['Tenant Admin', 'مدير النظام', 'Administrator inside each tenant company', 'Full access inside tenant DB'],
        ['Tenant Employee', 'موظف', 'Limited employee user (RBAC)', 'Limited by role & permissions'],
    ],
    col_widths=[Cm(3), Cm(3), Cm(5), Cm(5)],
)

add_heading(doc, '1.4 Project Status', level=2)
add_para(doc,
    'Per COMPLETION_SUMMARY_AR.md the system is at 83% stability rating. Major modules are working '
    'with full CRUD, validations, accounting posting, and Arabic RTL UI. The remaining work '
    'concentrates on: hardening tests, error handling on remaining services, caching, and rate-limiting '
    'sensitive operations.')

# ============================================================
# PART 2 — ARCHITECTURE
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 2: Application Architecture', level=1)

add_heading(doc, '2.1 Multi-Tenancy Architecture', level=2)
add_para(doc, 'Magzani uses Stancl Tenancy 3.x in database-per-tenant mode:')
features = [
    'Tenant model: App\\Models\\Tenant (extends Stancl\\Tenancy\\Database\\Models\\Tenant)',
    'ID generator: Stancl\\Tenancy\\UUIDGenerator',
    'Tenant DB naming: tenant{tenant_id} (prefix = "tenant", suffix = "")',
    'Active bootstrappers: DatabaseTenancyBootstrapper, FilesystemTenancyBootstrapper, QueueTenancyBootstrapper',
    'Cache & Redis bootstrappers are intentionally commented out',
    'Tenant migrations path: database/migrations/tenant/ (configured in config/tenancy.php)',
    'Central domains: 127.0.0.1, localhost, plus env CENTRAL_DOMAINS (comma-separated)',
    'Tenant domain pattern: {tenant_id}.{TENANT_DOMAIN_SUFFIX} (default suffix = localhost)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_heading(doc, '2.2 Landlord vs Tenant Separation', level=2)
add_para(doc, 'Two completely separate surfaces share the same Laravel installation:')

add_table(doc,
    ['Aspect', 'Landlord (Central)', 'Tenant (Per-DB)'],
    [
        ['Migrations', 'database/migrations/', 'database/migrations/tenant/'],
        ['Routes file', 'routes/web.php (prefix /super-admin)', 'routes/tenant.php (619 lines)'],
        ['Models', 'Tenant, Plan, PlanFeature, Domain', 'User, Product, Invoice, etc. (~74 models)'],
        ['Layout', 'resources/views/landlord/layout.blade.php', 'resources/views/layouts/app.blade.php'],
        ['Branding', 'Kayyan SaaS (K letter SVG)', 'Magzani ERP (RTL Arabic)'],
        ['Middleware order', 'web + auth + central-domain', 'web + PreventCentral + InitializeTenant + feature'],
        ['Auth provider', 'App\\Models\\Tenant (Stancl)', 'App\\Models\\User'],
    ],
    col_widths=[Cm(3.5), Cm(6.5), Cm(6.5)],
)

add_heading(doc, '2.3 Request Lifecycle', level=2)
add_para(doc, 'When a request hits the application:')
steps = [
    '1. Laravel boots, registers service providers (including TenancyServiceProvider and AccountingServiceProvider).',
    '2. AppServiceProvider registers TenantObserver on Tenant, and a view composer that injects $planFeatures into ALL views.',
    '3. RouteServiceProvider loads routes/web.php (landlord) and routes/tenant.php (tenant).',
    '4. Tenant routes group applies middleware: web → PreventAccessFromCentralDomains → InitializeTenancyByDomain → feature.',
    '5. InitializeTenancyByDomain reads the host, looks up the Domain row, finds the Tenant, and swaps the DB connection to tenant{id}.',
    '6. The `feature` middleware checks tenant()->hasFeature(key) against the central plans table.',
    '7. Controllers run inside the tenant DB context.',
    '8. On response, Stancl restores the central connection.',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '2.4 Tenancy Resolution Flow', level=2)
add_para(doc, 'Detailed lifecycle of tenancy bootstrap:')
flow = [
    'Request → tenant routes → middleware stack',
    '→ PreventAccessFromCentralDomains (rejects if host ∈ central domains)',
    '→ InitializeTenancyByDomain → Stancl resolves Domain → finds Tenant model',
    '→ DatabaseTenancyBootstrapper switches default connection to tenant DB',
    '→ FilesystemTenancyBootstrapper suffixes local/public paths (suffix_storage_path = true)',
    '→ QueueTenancyBootstrapper tags the current queue for tenant isolation',
    '→ feature middleware looks up tenant.plan_id (from JSON data) and validates feature flag',
    '→ Controller executes against tenant DB',
    '→ Response: connection restored',
]
for s in flow:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '2.5 Tenant Provisioning Flow', level=2)
add_para(doc, 'When super admin creates a new tenant from /super-admin/tenants (SuperAdminController::tenantsStore):')
prov_steps = [
    '1. Validate tenant_id field (alpha_dash + lowercase + unique).',
    '2. Create Tenant model — fires Events\\TenantCreated.',
    '3. Stancl Tenancy queues CreateDatabase job → creates tenant{id} database.',
    '4. Stancl queues MigrateDatabase job → runs all migrations under database/migrations/tenant/.',
    '5. Create Domain: {tenant_id}.{TENANT_DOMAIN_SUFFIX}.',
    '6. Inside tenant()->run(...) context:',
    '   - Run PermissionAndRoleSeeder (creates admin & employee roles, all permissions)',
    '   - Create user admin@{tenant_id}.com with password "password"',
    '   - Attach admin role to user',
]
for s in prov_steps:
    doc.add_paragraph(s, style='List Bullet')

# Save what we have so we can append part 2 in a second script run
doc.save(OUT_PATH)
print(f"Part 1 saved → {OUT_PATH}")
