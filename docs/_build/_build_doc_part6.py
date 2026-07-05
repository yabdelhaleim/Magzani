"""
Magzani ERP Doc - Part 5
Complete Routes Reference (Landlord + Tenant) - كل الراوتس
"""
import sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from _build_doc_part5 import (
    doc, add_heading, add_para, add_para_rtl, add_table, add_code_block,
    add_hr, add_page_break, OUT_PATH,
)
from docx.shared import Pt, Cm, RGBColor

# ============================================================
# PART 5 - ROUTES REFERENCE
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 5: Complete Routes Reference', level=1)
add_para(doc,
    'This part lists every route in routes/web.php (landlord) and routes/tenant.php. '
    'All tenant routes are wrapped in middleware: web + PreventAccessFromCentralDomains + '
    'InitializeTenancyByDomain + feature.'
)

# ============================================================
# LANDLORD ROUTES
# ============================================================
add_heading(doc, '5.1 Landlord Routes (routes/web.php)', level=2)
add_para(doc, 'File: routes/web.php (38 lines). All landlord routes are prefixed /super-admin and require central-domain context.')

add_table(doc,
    ['Method', 'URI', 'Name', 'Controller@Action'],
    [
        ['GET',   '/super-admin/dashboard',                         'super-admin.dashboard',     'SuperAdminController@dashboard'],
        ['GET',   '/super-admin/plans',                             'super-admin.plans.index',   'plansIndex'],
        ['GET',   '/super-admin/plans/create',                      'super-admin.plans.create',  'plansCreate'],
        ['POST',  '/super-admin/plans',                             'super-admin.plans.store',   'plansStore'],
        ['GET',   '/super-admin/plans/{plan}/edit',                 'super-admin.plans.edit',    'plansEdit'],
        ['PUT',   '/super-admin/plans/{plan}',                      'super-admin.plans.update',  'plansUpdate'],
        ['DELETE','/super-admin/plans/{plan}',                      'super-admin.plans.destroy', 'plansDestroy'],
        ['GET',   '/super-admin/tenants',                           'super-admin.tenants.index',  'tenantsIndex'],
        ['GET',   '/super-admin/tenants/create',                    'super-admin.tenants.create', 'tenantsCreate'],
        ['POST',  '/super-admin/tenants',                           'super-admin.tenants.store',  'tenantsStore'],
        ['GET',   '/super-admin/tenants/{id}/edit',                 'super-admin.tenants.edit',   'tenantsEdit'],
        ['PUT',   '/super-admin/tenants/{id}',                      'super-admin.tenants.update', 'tenantsUpdate'],
        ['POST',  '/super-admin/tenants/{id}/toggle-status',        'super-admin.tenants.toggle-status', 'tenantsToggleStatus'],
        ['DELETE','/super-admin/tenants/{id}',                      'super-admin.tenants.destroy', 'tenantsDestroy'],
    ],
    col_widths=[Cm(1.7), Cm(5.5), Cm(5), Cm(4)],
    font_size=9,
)

# ============================================================
# TENANT ROUTES
# ============================================================
add_page_break(doc)
add_heading(doc, '5.2 Tenant Routes (routes/tenant.php — 619 lines)', level=2)

add_para(doc, 'The full middleware stack wrapping all tenant routes:')
add_code_block(doc, """Route::middleware([
    'web',
    \\Stancl\\Tenancy\\Middleware\\PreventAccessFromCentralDomains::class,
    \\Stancl\\Tenancy\\Middleware\\InitializeTenancyByDomain::class,
    'feature',
])->group(function () {
    // ... all tenant routes
});""")

# ----- AUTH -----
add_heading(doc, '5.2.1 Auth & Public', level=3)
add_table(doc,
    ['Method', 'URI', 'Name', 'Middleware'],
    [
        ['GET',  '/login',          'login',     'guest'],
        ['POST', '/login',          'login.post','guest + throttle:5,1'],
        ['POST', '/logout',         'logout',    'auth'],
        ['GET',  '/',               'dashboard', 'auth + role middleware'],
        ['GET',  '/home',           'home',      'redirects → /'],
        ['GET',  '/plan/upgrade',   'plan.upgrade', 'PlanController'],
    ],
    col_widths=[Cm(1.5), Cm(5), Cm(5), Cm(5)],
    font_size=9,
)

# ----- POS -----
add_heading(doc, '5.2.2 POS & Shifts (feature:pos)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name', 'Notes'],
    [
        ['GET',  '/pos',                  'pos.index',        'PosPanel Livewire'],
        ['GET',  '/pos/returns',          'pos.returns',      'PosReturnPanel Livewire'],
        ['GET',  '/pos/settings',         'pos.settings.index',   'admin.only'],
        ['POST', '/pos/settings',         'pos.settings.update', 'admin.only'],
        ['GET',  '/pos/shift/open',       'pos.shift.create', 'open form'],
        ['POST', '/pos/shift/open',       'pos.shift.open',   'open action'],
        ['GET',  '/pos/shift/close',      'pos.shift.close-view', 'close form'],
        ['POST', '/pos/shift/close',      'pos.shift.close',  'close action'],
        ['GET',  '/pos/history',          'pos.history',      'past shifts'],
        ['GET',  '/pos/x-report',         'pos.xreport',      'mid-shift report'],
        ['GET',  '/pos/shift/{id}/z-report', 'pos.shift.zreport', 'end-of-day report'],
    ],
    col_widths=[Cm(1.5), Cm(5.5), Cm(4.5), Cm(5)],
    font_size=9,
)

# ----- SALES INVOICES (POS-flavoured) -----
add_heading(doc, '5.2.3 Sales Invoices (feature:pos, admin-only writes)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/invoices/sales',                       'invoices.sales.index'],
        ['GET',    '/invoices/sales/create',                'invoices.sales.create'],
        ['POST',   '/invoices/sales',                       'invoices.sales.store'],
        ['GET',    '/invoices/sales/{id}/edit',             'invoices.sales.edit'],
        ['PUT',    '/invoices/sales/{id}',                  'invoices.sales.update'],
        ['DELETE', '/invoices/sales/{id}',                  'invoices.sales.destroy'],
        ['GET',    '/invoices/sales/{id}',                  'invoices.sales.show'],
        ['GET',    '/invoices/sales/{id}/print',            'invoices.sales.print'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- MANUFACTURING -----
add_heading(doc, '5.2.4 Manufacturing (feature:manufacturing + admin.only)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/manufacturing',                            'manufacturing.index'],
        ['GET',    '/manufacturing/create',                     'manufacturing.create'],
        ['POST',   '/manufacturing',                            'manufacturing.store'],
        ['POST',   '/manufacturing/calculate',                  'manufacturing.calculate'],
        ['GET',    '/manufacturing/{id}',                       'manufacturing.show'],
        ['GET',    '/manufacturing/{id}/edit',                  'manufacturing.edit'],
        ['PUT',    '/manufacturing/{id}',                       'manufacturing.update'],
        ['DELETE', '/manufacturing/{id}',                       'manufacturing.destroy'],
        ['POST',   '/manufacturing/{id}/confirm',               'manufacturing.confirm'],
        # Manufacturing orders
        ['GET',    '/manufacturing-orders',                     'manufacturing-orders.index'],
        ['GET',    '/manufacturing-orders/create',              'manufacturing-orders.create'],
        ['POST',   '/manufacturing-orders',                     'manufacturing-orders.store'],
        ['POST',   '/manufacturing-orders/calculate',           'manufacturing-orders.calculate'],
        ['GET',    '/manufacturing-orders/{id}',                'manufacturing-orders.show'],
        ['GET',    '/manufacturing-orders/{id}/edit',           'manufacturing-orders.edit'],
        ['PUT',    '/manufacturing-orders/{id}',                'manufacturing-orders.update'],
        ['DELETE', '/manufacturing-orders/{id}',                'manufacturing-orders.destroy'],
        ['POST',   '/manufacturing-orders/{id}/confirm',        'manufacturing-orders.confirm'],
        ['PATCH',  '/manufacturing-orders/{id}/complete',       'manufacturing-orders.complete'],
        ['PATCH',  '/manufacturing-orders/{id}/cancel',         'manufacturing-orders.cancel'],
        # Wood stocks + dispensings
        ['GET',    '/manufacturing/wood-stocks',                'manufacturing.wood-stocks.index'],
        ['GET',    '/manufacturing/wood-stocks/create',         'manufacturing.wood-stocks.create'],
        ['POST',   '/manufacturing/wood-stocks',                'manufacturing.wood-stocks.store'],
        ['GET',    '/manufacturing/wood-stocks/{id}/edit',      'manufacturing.wood-stocks.edit'],
        ['PUT',    '/manufacturing/wood-stocks/{id}',           'manufacturing.wood-stocks.update'],
        ['DELETE', '/manufacturing/wood-stocks/{id}',           'manufacturing.wood-stocks.destroy'],
        ['GET',    '/manufacturing/wood-dispensings',           'manufacturing.wood-dispensings.index'],
        ['GET',    '/manufacturing/wood-dispensings/create',    'manufacturing.wood-dispensings.create'],
        ['POST',   '/manufacturing/wood-dispensings',           'manufacturing.wood-dispensings.store'],
        ['GET',    '/manufacturing/wood-dispensings/{id}/edit', 'manufacturing.wood-dispensings.edit'],
        ['PUT',    '/manufacturing/wood-dispensings/{id}',      'manufacturing.wood-dispensings.update'],
        ['DELETE', '/manufacturing/wood-dispensings/{id}',      'manufacturing.wood-dispensings.destroy'],
        # Raw materials sub-resource
        ['GET',    '/manufacturing-orders/raw-materials',             'manufacturing-orders.raw-materials.index'],
        ['GET',    '/manufacturing-orders/raw-materials/create',      'manufacturing-orders.raw-materials.create'],
        ['POST',   '/manufacturing-orders/raw-materials',             'manufacturing-orders.raw-materials.store'],
        ['GET',    '/manufacturing-orders/raw-materials/{id}/edit',   'manufacturing-orders.raw-materials.edit'],
        ['PUT',    '/manufacturing-orders/raw-materials/{id}',        'manufacturing-orders.raw-materials.update'],
        ['DELETE', '/manufacturing-orders/raw-materials/{id}',        'manufacturing-orders.raw-materials.destroy'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- WAREHOUSE TRANSFERS -----
add_heading(doc, '5.2.5 Warehouse Transfers (feature:multi_warehouse)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/transfers',                            'transfers.index'],
        ['GET',    '/transfers/create',                     'transfers.create'],
        ['POST',   '/transfers',                            'transfers.store'],
        ['GET',    '/transfers/pending',                    'transfers.pending'],
        ['GET',    '/transfers/warehouse/{id}/history',     'transfers.warehouse-history'],
        ['GET',    '/transfers/{id}',                       'transfers.show'],
        ['POST',   '/transfers/{id}/reverse',               'transfers.reverse'],
        ['POST',   '/transfers/{id}/cancel',                'transfers.cancel'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- ACCOUNTING QUICK + REPORTS -----
add_heading(doc, '5.2.6 Accounting (Quick) & Reports (feature:accounting, admin)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',  '/accounting/treasury',                    'accounting.treasury'],
        ['GET',  '/accounting/payments',                    'accounting.payments'],
        ['POST', '/accounting/payments/deposits',           'accounting.payments.deposits'],
        ['POST', '/accounting/payments/withdrawals',        'accounting.payments.withdrawals'],
        ['PUT',  '/accounting/payments/transactions/{id}',  'accounting.payments.transactions.update'],
        ['DELETE','/accounting/payments/transactions/{id}', 'accounting.payments.transactions.destroy'],
        ['GET',  '/accounting/expenses',                    'accounting.expenses'],
        ['POST', '/accounting/expenses',                    'accounting.expenses.store'],
        ['GET',  '/accounting/expenses/statistics',         'accounting.expenses.statistics'],
        # Reports
        ['GET',  '/reports/financial',                      'reports.financial'],
        ['GET',  '/reports/inventory',                      'reports.inventory'],
        ['GET',  '/reports/profit-loss',                    'reports.profit-loss'],
        ['GET',  '/reports/financial/export',               'reports.financial.export'],
        ['GET',  '/reports/inventory/export',               'reports.inventory.export'],
        ['GET',  '/reports/profit-loss/export',             'reports.profit-loss.export'],
        ['GET',  '/reports/wood-stock',                     'reports.wood-stock'],
        ['GET',  '/reports/wood-movement',                  'reports.wood-movement'],
        ['GET',  '/reports/wood-cost-production',           'reports.wood-cost-production'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- WAREHOUSES -----
add_heading(doc, '5.2.7 Warehouses (feature:warehouses + role)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/warehouses',                       'warehouses.index'],
        ['GET',    '/warehouses/create',                'warehouses.create'],
        ['POST',   '/warehouses',                       'warehouses.store'],
        ['GET',    '/warehouses/{id}',                  'warehouses.show'],
        ['GET',    '/warehouses/{id}/edit',             'warehouses.edit'],
        ['PUT',    '/warehouses/{id}',                  'warehouses.update'],
        ['DELETE', '/warehouses/{id}',                  'warehouses.destroy'],
        ['POST',   '/warehouses/{id}/add-product',      'warehouses.add-product'],
        ['GET',    '/warehouses/low-stock',             'warehouses.low-stock'],
        ['GET',    '/warehouses/{id}/movements',        'warehouses.movements'],
        ['GET',    '/warehouses/search',                'warehouses.search'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- WAREHOUSE ORDERS -----
add_heading(doc, '5.2.8 Warehouse Orders (inbound + outbound)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',  '/warehouse-orders/stock-preview',                         'warehouse-orders.stock-preview'],
        ['GET',  '/warehouse-orders/inbound',                                'warehouse-orders.inbound.index'],
        ['GET',  '/warehouse-orders/inbound/create',                         'warehouse-orders.inbound.create'],
        ['POST', '/warehouse-orders/inbound',                                'warehouse-orders.inbound.store'],
        ['GET',  '/warehouse-orders/inbound/{id}',                           'warehouse-orders.inbound.show'],
        ['GET',  '/warehouse-orders/inbound/{id}/print',                     'warehouse-orders.inbound.print'],
        ['GET',  '/warehouse-orders/outbound',                               'warehouse-orders.outbound.index'],
        ['GET',  '/warehouse-orders/outbound/create',                        'warehouse-orders.outbound.create'],
        ['POST', '/warehouse-orders/outbound',                               'warehouse-orders.outbound.store'],
        ['GET',  '/warehouse-orders/outbound/{id}',                          'warehouse-orders.outbound.show'],
        ['GET',  '/warehouse-orders/outbound/{id}/print',                    'warehouse-orders.outbound.print'],
        ['POST', '/warehouse-orders/outbound/{id}/approve',                  'warehouse-orders.outbound.approve'],
        ['POST', '/warehouse-orders/outbound/{id}/cancel',                   'warehouse-orders.outbound.cancel'],
    ],
    col_widths=[Cm(1.5), Cm(8.5), Cm(6.5)],
    font_size=9,
)

# ----- STOCK COUNTS -----
add_heading(doc, '5.2.9 Stock Counts', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/stock-counts',                        'stock-counts.index'],
        ['GET',    '/stock-counts/create',                 'stock-counts.create'],
        ['POST',   '/stock-counts',                        'stock-counts.store'],
        ['GET',    '/stock-counts/{id}',                   'stock-counts.show'],
        ['POST',   '/stock-counts/{id}/start',             'stock-counts.start'],
        ['POST',   '/stock-counts/{id}/count',             'stock-counts.count'],
        ['POST',   '/stock-counts/{id}/complete',          'stock-counts.complete'],
        ['POST',   '/stock-counts/{id}/cancel',            'stock-counts.cancel'],
        ['POST',   '/stock-counts/{id}/items/{itemId}/approve', 'stock-counts.items.approve'],
        ['POST',   '/stock-counts/{id}/items/{itemId}/count',   'stock-counts.items.count'],
        ['POST',   '/stock-counts/{id}/approve-all',       'stock-counts.approve-all'],
        ['GET',    '/stock-counts/warehouse/{id}/products', 'stock-counts.warehouse-products'],
        ['GET',    '/stock-counts/{id}/print',             'stock-counts.print'],
    ],
    col_widths=[Cm(1.5), Cm(7), Cm(7)],
    font_size=9,
)

# ----- INVENTORY MOVEMENTS -----
add_heading(doc, '5.2.10 Inventory Movements', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET', '/movements',                    'movements.index'],
        ['GET', '/movements/product/{id}',       'movements.product'],
        ['GET', '/movements/export',             'movements.export'],
    ],
    col_widths=[Cm(1.7), Cm(7), Cm(7)],
    font_size=9,
)

# ----- PRODUCTS -----
add_heading(doc, '5.2.11 Products (auth)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name', 'Notes'],
    [
        ['GET',    '/products',                            'products.index', ''],
        ['GET',    '/products/barcode/print',              'products.barcode.print', ''],
        ['GET',    '/products/bulk-price-update',          'products.bulk-price-update.index', ''],
        ['POST',   '/products/bulk-price-update/preview',  'products.bulk-price-update.preview', ''],
        ['POST',   '/products/bulk-price-update/apply',    'products.bulk-price-update.apply', ''],
        ['POST',   '/products/categories-by-unit',         'products.categories-by-unit', 'AJAX'],
        ['POST',   '/products/by-unit-category',           'products.by-unit-category', 'AJAX'],
        ['POST',   '/products/preview-smart-update',       'products.preview-smart-update', 'AJAX'],
        ['POST',   '/products/suggested-pricing',          'products.suggested-pricing', 'AJAX'],
        ['POST',   '/products/convert-unit-price',         'products.convert-unit-price', 'AJAX'],
        ['POST',   '/products/unit-details',               'products.unit-details', 'AJAX'],
        ['GET',    '/products/create',                     'products.create', 'admin.only'],
        ['POST',   '/products',                            'products.store', 'admin.only'],
        ['GET',    '/products/{product}/edit',             'products.edit', 'admin.only'],
        ['PUT',    '/products/{product}',                  'products.update', 'admin.only'],
        ['DELETE', '/products/{product}',                  'products.destroy', 'admin.only'],
        ['GET',    '/products/{product}',                  'products.show', ''],
        ['POST',   '/products/{product}/update-price',     'products.update-price', ''],
        ['GET',    '/products/{product}/price-history',    'products.price-history', ''],
    ],
    col_widths=[Cm(1.4), Cm(6.8), Cm(5), Cm(2.5)],
    font_size=9,
)

# ----- CATEGORIES -----
add_heading(doc, '5.2.12 Categories (auth)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/categories',                  'categories.index'],
        ['POST',   '/categories',                  'categories.store'],
        ['PUT',    '/categories/{id}',             'categories.update'],
        ['DELETE', '/categories/{id}',             'categories.destroy'],
        ['POST',   '/categories/{id}/toggle-status','categories.toggle-status'],
        ['GET',    '/categories/list',             'categories.list'],
    ],
    col_widths=[Cm(1.5), Cm(6.5), Cm(7)],
    font_size=9,
)

# ----- PURCHASES -----
add_heading(doc, '5.2.13 Purchases (feature:purchase + auth)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/invoices/purchases',                        'invoices.purchases.index'],
        ['GET',    '/invoices/purchases/create',                 'invoices.purchases.create'],
        ['POST',   '/invoices/purchases',                        'invoices.purchases.store'],
        ['GET',    '/invoices/purchases/{id}',                   'invoices.purchases.show'],
        ['GET',    '/invoices/purchases/{id}/edit',              'invoices.purchases.edit'],
        ['PUT',    '/invoices/purchases/{id}',                   'invoices.purchases.update'],
        ['DELETE', '/invoices/purchases/{id}',                   'invoices.purchases.destroy'],
        ['POST',   '/invoices/purchases/{id}/confirm',           'invoices.purchases.confirm'],
        ['GET',    '/invoices/purchases/{id}/print',             'invoices.purchases.print'],
        ['GET',    '/invoices/purchases/export',                 'invoices.purchases.export'],
        ['GET',    '/invoices/purchases/{id}/exportSingle',      'invoices.purchases.exportSingle'],
        ['GET',    '/invoices/purchase-returns',                 'invoices.purchase-returns.index'],
        ['GET',    '/invoices/purchase-returns/create',          'invoices.purchase-returns.create'],
        ['POST',   '/invoices/purchase-returns',                 'invoices.purchase-returns.store'],
        ['GET',    '/invoices/purchase-returns/{id}',            'invoices.purchase-returns.show'],
        ['GET',    '/invoices/purchase-returns/{id}/edit',       'invoices.purchase-returns.edit'],
        ['PUT',    '/invoices/purchase-returns/{id}',            'invoices.purchase-returns.update'],
        ['DELETE', '/invoices/purchase-returns/{id}',            'invoices.purchase-returns.destroy'],
        ['POST',   '/invoices/purchase-returns/available-items', 'invoices.purchase-returns.available-items'],
    ],
    col_widths=[Cm(1.5), Cm(7.5), Cm(7)],
    font_size=9,
)

# ----- SALES RETURNS -----
add_heading(doc, '5.2.14 Sales Returns (feature:pos + auth)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/invoices/sales-returns',               'invoices.sales-returns.index'],
        ['GET',    '/invoices/sales-returns/create',        'invoices.sales-returns.create'],
        ['POST',   '/invoices/sales-returns',               'invoices.sales-returns.store'],
        ['GET',    '/invoices/sales-returns/{id}',          'invoices.sales-returns.show'],
        ['GET',    '/invoices/sales-returns/{id}/edit',     'invoices.sales-returns.edit'],
        ['PUT',    '/invoices/sales-returns/{id}',          'invoices.sales-returns.update'],
        ['DELETE', '/invoices/sales-returns/{id}',          'invoices.sales-returns.destroy'],
    ],
    col_widths=[Cm(1.5), Cm(7.5), Cm(7)],
    font_size=9,
)

# ----- CUSTOMERS/SUPPLIERS -----
add_heading(doc, '5.2.15 Customers & Suppliers (auth)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/customers',                          'customers.index'],
        ['GET',    '/customers/create',                   'customers.create'],
        ['POST',   '/customers',                          'customers.store'],
        ['GET',    '/customers/{id}',                     'customers.show'],
        ['GET',    '/customers/{id}/edit',                'customers.edit'],
        ['PUT',    '/customers/{id}',                     'customers.update'],
        ['DELETE', '/customers/{id}',                     'customers.destroy'],
        ['GET',    '/customers/{id}/statement',           'customers.statement'],
        ['GET',    '/customers/{id}/statement/export',    'customers.statement.export'],
        ['GET',    '/suppliers',                          'suppliers.index'],
        ['GET',    '/suppliers/create',                   'suppliers.create'],
        ['POST',   '/suppliers',                          'suppliers.store'],
        ['GET',    '/suppliers/{id}',                     'suppliers.show'],
        ['GET',    '/suppliers/{id}/edit',                'suppliers.edit'],
        ['PUT',    '/suppliers/{id}',                     'suppliers.update'],
        ['DELETE', '/suppliers/{id}',                     'suppliers.destroy'],
        ['GET',    '/suppliers/{id}/statement',           'suppliers.statement'],
        ['GET',    '/suppliers/{id}/statement/export',    'suppliers.statement.export'],
    ],
    col_widths=[Cm(1.5), Cm(7), Cm(7)],
    font_size=9,
)

# ----- SETTINGS -----
add_heading(doc, '5.2.16 Settings (auth + admin)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',   '/settings',                    'settings.index'],
        ['POST',  '/settings/company',            'settings.company'],
        ['POST',  '/settings/system',             'settings.system'],
        ['POST',  '/settings/logo/delete',        'settings.logo.delete'],
    ],
    col_widths=[Cm(1.5), Cm(7), Cm(7)],
    font_size=9,
)

# ----- USERS / PERMISSIONS -----
add_heading(doc, '5.2.17 Users & Permissions (auth + admin)', level=3)
add_table(doc,
    ['Method', 'URI', 'Name'],
    [
        ['GET',    '/users',                            'users.index'],
        ['GET',    '/users/create',                     'users.create'],
        ['POST',   '/users',                            'users.store'],
        ['GET',    '/users/{user}',                     'users.show'],
        ['GET',    '/users/{user}/edit',                'users.edit'],
        ['PUT',    '/users/{user}',                     'users.update'],
        ['DELETE', '/users/{user}',                     'users.destroy'],
        ['POST',   '/users/{user}/toggle-active',       'users.toggle-active'],
        ['GET',    '/permissions',                      'permissions.index'],
        ['GET',    '/permissions/print',                'permissions.print'],
        ['GET',    '/permissions/users/{user}/edit',    'permissions.users.edit'],
        ['PUT',    '/permissions/users/{user}',         'permissions.users.update'],
        ['GET',    '/permissions/roles',                'permissions.roles.index'],
        ['POST',   '/permissions/roles',                'permissions.roles.store'],
        ['GET',    '/permissions/roles/{role}/edit',    'permissions.roles.edit'],
        ['PUT',    '/permissions/roles/{role}',         'permissions.roles.update'],
        ['DELETE', '/permissions/roles/{role}',         'permissions.roles.destroy'],
        ['PUT',    '/permissions/roles/{role}/permissions', 'permissions.roles.update-permissions'],
    ],
    col_widths=[Cm(1.5), Cm(7), Cm(7)],
    font_size=9,
)

# ----- ADVANCED ACCOUNTING -----
add_heading(doc, '5.2.18 Advanced Accounting (feature:accounting_advanced + admin)', level=3)
acc_routes = [
    ['GET',   '/accounting',                                          'accounting.dashboard'],
    ['GET',   '/accounting/integrity-check',                         'accounting.integrity-check'],
    ['POST',  '/accounting/integrity-check/fix',                      'accounting.integrity-check.fix'],
    # Setup wizard
    ['GET',   '/accounting/setup',                                    'accounting.setup.index'],
    # COA
    ['GET',   '/accounting/coa',                                      'accounting.coa.index'],
    ['GET',   '/accounting/coa/create',                               'accounting.coa.create'],
    ['POST',  '/accounting/coa',                                      'accounting.coa.store'],
    ['GET',   '/accounting/coa/{id}/edit',                            'accounting.coa.edit'],
    ['PUT',   '/accounting/coa/{id}',                                 'accounting.coa.update'],
    ['DELETE','/accounting/coa/{id}',                                 'accounting.coa.destroy'],
    ['GET',   '/accounting/coa/export',                               'accounting.coa.export'],
    # Journal entries
    ['GET',   '/accounting/journal',                                  'accounting.journal.index'],
    ['GET',   '/accounting/journal/create',                           'accounting.journal.create'],
    ['POST',  '/accounting/journal',                                  'accounting.journal.store'],
    ['GET',   '/accounting/journal/{id}',                             'accounting.journal.show'],
    ['PUT',   '/accounting/journal/{id}',                             'accounting.journal.update'],
    ['DELETE','/accounting/journal/{id}',                             'accounting.journal.destroy'],
    ['POST',  '/accounting/journal/{id}/post',                        'accounting.journal.post'],
    ['POST',  '/accounting/journal/{id}/reverse',                     'accounting.journal.reverse'],
    # Vouchers
    ['GET',   '/accounting/vouchers/receipt',                         'accounting.vouchers.receipt.index'],
    ['POST',  '/accounting/vouchers/receipt',                         'accounting.vouchers.receipt.store'],
    ['GET',   '/accounting/vouchers/receipt/{id}/print',              'accounting.vouchers.receipt.print'],
    ['GET',   '/accounting/vouchers/payment',                         'accounting.vouchers.payment.index'],
    ['POST',  '/accounting/vouchers/payment',                         'accounting.vouchers.payment.store'],
    ['GET',   '/accounting/vouchers/payment/{id}/print',              'accounting.vouchers.payment.print'],
    # Reports
    ['GET',   '/accounting/reports/trial-balance',                    'accounting.reports.trial-balance'],
    ['GET',   '/accounting/reports/income-statement',                 'accounting.reports.income-statement'],
    ['GET',   '/accounting/reports/comparative-income',               'accounting.reports.comparative-income'],
    ['GET',   '/accounting/reports/balance-sheet',                    'accounting.reports.balance-sheet'],
    ['GET',   '/accounting/reports/general-ledger',                   'accounting.reports.general-ledger'],
    ['GET',   '/accounting/reports/partner-ledger',                   'accounting.reports.partner-ledger'],
    ['GET',   '/accounting/reports/audit-trail',                      'accounting.reports.audit-trail'],
    ['GET',   '/accounting/reports/aging',                            'accounting.reports.aging'],
    ['GET',   '/accounting/reports/financial-ratios',                 'accounting.reports.financial-ratios'],
    ['GET',   '/accounting/reports/vat-settlement',                   'accounting.reports.vat-settlement'],
    # Recurring
    ['GET',   '/accounting/recurring',                                'accounting.recurring.index'],
    ['POST',  '/accounting/recurring',                                'accounting.recurring.store'],
    ['GET',   '/accounting/recurring/{id}/edit',                      'accounting.recurring.edit'],
    ['PUT',   '/accounting/recurring/{id}',                           'accounting.recurring.update'],
    ['DELETE','/accounting/recurring/{id}',                           'accounting.recurring.destroy'],
    ['POST',  '/accounting/recurring/{id}/run-now',                   'accounting.recurring.run-now'],
    # Fiscal periods
    ['GET',   '/accounting/fiscal',                                  'accounting.fiscal.index'],
    ['GET',   '/accounting/fiscal/years/create',                      'accounting.fiscal.years.create'],
    ['POST',  '/accounting/fiscal/years',                             'accounting.fiscal.years.store'],
    ['POST',  '/accounting/fiscal/years/{year}/close',                'accounting.fiscal.years.close'],
    ['POST',  '/accounting/fiscal/years/{year}/reopen',               'accounting.fiscal.years.reopen'],
    ['POST',  '/accounting/fiscal/years/{year}/year-end-close',       'accounting.fiscal.years.year-end-close'],
    # Posting failures
    ['GET',   '/accounting/posting-failures',                         'accounting.posting-failures.index'],
    ['POST',  '/accounting/posting-failures/{failure}/retry',         'accounting.posting-failures.retry'],
    ['POST',  '/accounting/posting-failures/{failure}/resolve',       'accounting.posting-failures.resolve'],
    # Fixed assets
    ['GET',   '/accounting/fixed-assets',                             'accounting.fixed-assets.index'],
    ['POST',  '/accounting/fixed-assets',                             'accounting.fixed-assets.store'],
    ['GET',   '/accounting/fixed-assets/{id}/edit',                   'accounting.fixed-assets.edit'],
    ['PUT',   '/accounting/fixed-assets/{id}',                        'accounting.fixed-assets.update'],
    ['DELETE','/accounting/fixed-assets/{id}',                        'accounting.fixed-assets.destroy'],
    ['POST',  '/accounting/fixed-assets/{id}/depreciate',             'accounting.fixed-assets.depreciate'],
    ['POST',  '/accounting/fixed-assets/{id}/dispose',                'accounting.fixed-assets.dispose'],
    # Settings
    ['GET',   '/accounting/settings',                                 'accounting.settings.index'],
    ['POST',  '/accounting/settings',                                 'accounting.settings.update'],
]
add_table(doc,
    ['Method', 'URI', 'Name'],
    acc_routes,
    col_widths=[Cm(1.5), Cm(7.5), Cm(7)],
    font_size=9,
)

# ----- API -----
add_heading(doc, '5.3 API Routes (routes/api.php)', level=2)
add_para(doc, 'Only ONE route is exposed via API right now:')
add_table(doc,
    ['Method', 'URI', 'Name', 'Middleware', 'Returns'],
    [
        ['GET', '/api/user', 'api.user', 'auth:sanctum', '$request->user()'],
    ],
    col_widths=[Cm(1.7), Cm(3), Cm(3.5), Cm(3), Cm(4)],
    font_size=9,
)

doc.save(OUT_PATH)
print(f"Part 5 (Routes) appended → {OUT_PATH}")
