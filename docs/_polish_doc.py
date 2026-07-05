"""
Final polish:
1. Reduce cover title from 48pt to 44pt
2. Consolidate empty paragraphs on cover
3. Add updateFields=true setting so Word prompts to refresh TOC on open
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
doc = Document(IN_PATH)

# --- Fix 1: Reduce cover title font ---
# Find the "Magzani ERP" heading (cover title)
# This is the first paragraph that contains "Magzani ERP"
title_found = False
for p in doc.paragraphs[:10]:
    if 'Magzani ERP' in p.text and not title_found:
        for run in p.runs:
            if 'Magzani ERP' in run.text:
                run.font.size = Pt(44)
                print(f"Reduced cover title to 44pt: '{run.text}'")
                title_found = True
                break

# --- Fix 2: Reduce consecutive empty paragraphs on cover ---
# Look at first 10 paragraphs, remove excess blank ones
empties = 0
removed = 0
for i, p in enumerate(doc.paragraphs[:15]):
    if not p.text.strip():
        empties += 1
        if empties > 2:
            # Could remove, but safer to leave as-is — adjust spacing instead
            pass
    else:
        empties = 0

# A simpler safer fix: reduce space_before/after on empty paragraphs at top
for i, p in enumerate(doc.paragraphs[:8]):
    if not p.text.strip():
        # Reduce space before/after to collapse visually
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

# --- Fix 3: Add updateFields=true to settings.xml ---
settings = doc.settings.element
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Check if updateFields exists
update_fields = settings.find(qn('w:updateFields'))
if update_fields is None:
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    settings.append(uf)
    print('Added w:updateFields=true to settings.xml')
else:
    update_fields.set(qn('w:val'), 'true')
    print('Set w:updateFields=true')

doc.save(IN_PATH)
print(f'\nSaved: {IN_PATH}')
print(f'Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB')
