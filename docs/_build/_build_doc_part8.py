"""
Magzani ERP Doc - Part 7 + Part 8
Recent Changes & Known Issues + Quick Reference Card
"""
import sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from _build_doc_part7 import (
    doc, add_heading, add_para, add_para_rtl, add_table, add_code_block,
    add_hr, add_page_break, OUT_PATH,
)
from docx.shared import Pt, Cm, RGBColor

# ============================================================
# PART 7 - RECENT CHANGES & KNOWN ISSUES
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 7: Recent Changes & Known Issues', level=1)

# ----- Recent Git -----
add_heading(doc, '7.1 Recent Git History (git log --oneline -15)', level=2)
recent_commits = [
    ['f60b6b8', 'CHANG *****', 'top of main'],
    ['10637c0', 'CHANG ***', 'recent refactor'],
    ['1d220ad', 'CHANG **',  'recent refactor'],
    ['045b0eb', 'CHANG *',   'recent refactor'],
    ['8e7e8e6', 'CHANG *',   'recent refactor'],
    ['d14c1a9', 'Merge PR #1 from yabdelhaleim/fix/route-500-errors', 'PR merge'],
    ['ef57130', 'fix: add missing controller methods and views for all routes returning 500', 'BUG FIX'],
    ['bd22115', 'feat: raw materials management system for manufacturing orders', 'FEAT'],
    ['5b42b62', 'magzaniv6', 'release/snapshot'],
    ['5d3a409', 'fix: mobile responsive design for settings and permissions pages', 'UI FIX'],
    ['ba001f6', 'fix: null-safe company access in reports, pass missing variables to views', 'BUG FIX'],
    ['b12e8d6', 'fix: null-safe company access in sales invoice print and add ProductService::updateStock', 'BUG FIX'],
    ['3f52e9d', 'fix: route ordering bug in products and UI fixes for customers/suppliers', 'BUG FIX'],
    ['fdd0a48', 'Fix show page buttons contrast and improve mobile responsiveness', 'UI FIX'],
    ['60e4711', 'Fix supplier statement print button contrast', 'UI FIX'],
    ['ec3b50c', 'Add null-safe checks in customer and supplier statement views', 'BUG FIX'],
    ['8732779', 'Fix route ordering for customers and suppliers to prevent /create matching as ID', 'BUG FIX'],
    ['ef6f60b', 'Unify manufacturing orders form: single responsive table with mobile card-view CSS', 'UI REFACTOR'],
    ['1b5fcca', 'Remove duplicate permissions migration, add warehouse order migrations', 'MIGRATION'],
    ['b5825d6', 'Fix wallet currency based on country', 'FIX'],
    ['d934086', 'Add manufacturing cost calculator (BOM-based pricing module)', 'FEAT'],
]
add_table(doc,
    ['Commit', 'Description', 'Type'],
    recent_commits,
    col_widths=[Cm(2), Cm(12), Cm(2)],
    font_size=8,
)

# ----- Recent File Activity -----
add_heading(doc, '7.2 Recent File Modifications (last 48h)', level=2)
add_para(doc, 'Files heavily modified on July 3-4, 2026 (most recent work):')
add_table(doc,
    ['File / Folder', 'What Likely Changed'],
    [
        ['app/Models/ActivityLog.php', 'Tweaks for activity tracking'],
        ['app/Models/Tenant.php', 'Multi-tenancy helpers updated'],
        ['app/Models/Supplier.php / Customer.php', 'Balance handling updates'],
        ['app/Models/PosShift.php', 'Cash diff & recalc methods'],
        ['app/Models/FixedAsset.php', 'Depreciation lifecycle'],
        ['app/Models/RecurringJournalEntry*.php', 'Recurring templates'],
        ['app/Models/AccountingSetting.php', 'New accounting toggles'],
        ['app/Models/CashTransaction.php', 'Treasury ledger'],
        ['app/Http/Controllers/PosShiftController.php', 'Lifecycle methods'],
        ['app/Http/Controllers/UserController.php', 'RBAC improvements'],
        ['app/Http/Controllers/Accounting/*', 'Multiple accounting controllers touched'],
        ['app/Providers/AppServiceProvider.php', 'View composer / observer'],
        ['app/Providers/AuthServiceProvider.php', 'Gates additions'],
        ['app/Providers/EventServiceProvider.php', 'Listeners registered'],
        ['app/Providers/AccountingServiceProvider.php', 'Service registration'],
        ['app/Services/*', 'Bulk updates across services'],
        ['database/seeders/{AccountingSettingsSeeder, DefaultChartOfAccountsSeeder, PlanSeeder}.php', 'Plan / COA / settings seeding refinements'],
        ['database/migrations/tenant/2026_06_*.php', 'New tenant migrations'],
        ['routes/tenant.php', '+2,156 bytes — major route additions'],
        ['config/tenancy.php', 'Multi-tenant config tweaks'],
        ['config/app.php', 'Provider registrations'],
        ['bootstrap/cache/services.php', 'Service cache invalidation'],
        ['database/migrations/tenant/2026_04_07_000001_update_customers_balance_column.php', 'Customer balance column update'],
        ['.env.example / .env.production.example', 'Env samples updated'],
    ],
    col_widths=[Cm(7), Cm(9)],
    font_size=8,
)

# ----- Known Issues -----
add_heading(doc, '7.3 Known Issues (from ERROR_FIXES_TODO.md, SYSTEM_AUDIT_REPORT.md)', level=2)

add_heading(doc, '7.3.1 Critical', level=3)
add_table(doc,
    ['Issue', 'Status'],
    [
        ['APP_DEBUG=true in .env → SECURITY RISK for production. Must be false.', 'Open — disable before deploy'],
        ['Missing try-catch in CustomerService::update / delete / updateBalance (likewise SupplierService, ProductService).', 'Open — partial fix done in recent commits'],
        ['Missing unique indexes on sales_invoices.invoice_number and purchase_invoices.invoice_number.', 'Open'],
        ['Missing foreign key constraints preventing deletes of customers/suppliers with invoices.', 'Open'],
        ['No rate limiting outside login (5/min) and journal post (30/min). Sensitive routes unprotected.', 'Open'],
    ],
    col_widths=[Cm(11), Cm(5)],
    font_size=9,
)

add_heading(doc, '7.3.2 Bugs Already Fixed (mention for awareness)', level=3)
add_table(doc,
    ['Bug', 'Fix'],
    [
        ['Routes 500 error on missing controller methods', 'Added missing controller methods + views (commit ef57130)'],
        ['Missing /create matching as ID in customers / suppliers / products', 'Route ordering bug fix (commits 3f52e9d, 8732779)'],
        ['Null-safe Company access in reports + sales invoice print', 'Passes missing variables, guards null (ba001f6, b12e8d6)'],
        ['Array access bug in customer & supplier statement', 'Use $transaction["type"] not $transaction->type (ec3b50c)'],
        ['Duplicate permissions migration', 'Removed (1b5fcca); warehouse order migrations added'],
        ['Mobile responsive design issues on settings/permissions', 'CSS adjustments (5d3a409)'],
        ['Mobile responsive customer statement', 'CSS adjustments (ba001f6)'],
        ['Button contrast on white backgrounds', 'White button text on transparent bg fixed (fdd0a48, 60e4711)'],
        ['Missing ProductService::updateStock', 'Added (b12e8d6)'],
        ['Wallet currency hardcoded', 'Now uses tenant country (b5825d6)'],
    ],
    col_widths=[Cm(7), Cm(9)],
    font_size=9,
)

add_heading(doc, '7.3.3 Module Stability (per COMPLETION_SUMMARY_AR.md)', level=3)
add_table(doc,
    ['Module', 'Stability', 'Notes'],
    [
        ['Invoices / Returns',     '90% — excellent'],
        ['Manufacturing',          '85% — very good'],
        ['Customers / Suppliers',  '85% — very good'],
        ['Inventory',              '80% — good'],
        ['Reports',                '80% — good'],
        ['Security',               '70% — needs hardening'],
        ['Performance',            '80% — good'],
        ['Overall',                '83% — very good'],
    ],
    col_widths=[Cm(4), Cm(3), Cm(9)],
    font_size=9,
)

add_heading(doc, '7.3.4 Operational Caveats', level=3)
caveats = [
    'No compiled Vue components are in the public/build folder at the moment — UI relies on Alpine + Livewire + CDN Tailwind. Vite is configured but not actively used in production.',
    'API routes are minimal: only /api/user is implemented. No documented REST API surface for tenants.',
    'activity_logs exists in TWO forms (legacy custom + Laravel default). Pick one in a future cleanup.',
    'The main README.md is the default Laravel boilerplate — never modified to describe the actual project (see DOCUMENTATION.md for real description).',
    'welcome.blade.php (27 KB) and welcome_central.blade.php (25 KB) are present but unused in normal flow.',
    'Reverb is installed but no broadcasting channels are defined (routes/channels.php has only 163 bytes of defaults).',
    'Controllers WarehouseOrderController (14 KB) and ReportingController (7.8 KB) contain significant logic — extract to services in a future refactor.',
    'The desktop project\\.php style plan structure (large file, .bak) suggests on-disk snapshot artifacts — DO NOT commit `.bak` files to Git.',
]
for c in caveats:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# PART 8 - QUICK REFERENCE CARD
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 8: Quick Reference Card', level=1)

add_heading(doc, '8.1 Default Credentials (per Plan Provisioning)', level=2)
add_table(doc,
    ['Surface', 'Email', 'Password', 'Where Set'],
    [
        ['Tenant Admin (per tenant)', 'admin@{tenant_id}.com', 'password', 'SuperAdminController::tenantsStore line 169'],
        ['Landlord Super Admin',     '(no default — manually create in DB or via tinker)', '-', 'Created by tenant itself'],
    ],
    col_widths=[Cm(5), Cm(5), Cm(2.5), Cm(4)],
)

add_heading(doc, '8.2 Important URLs (Local Dev)', level=2)
add_table(doc,
    ['URL', 'Surface'],
    [
        ['http://localhost/super-admin',                 'Landlord dashboard (central)'],
        ['http://localhost/super-admin/tenants',         'Manage tenants'],
        ['http://localhost/super-admin/tenants/create',  'Create new tenant'],
        ['http://localhost/super-admin/plans',           'Manage plans'],
        ['http://{tenant_id}.localhost',                 'Tenant login (e.g. http://acme.localhost)'],
        ['http://{tenant_id}.localhost/login',           'Login form'],
        ['http://{tenant_id}.localhost/pos',             'POS panel (if feature:pos enabled)'],
        ['http://{tenant_id}.localhost/dashboard',       'Dashboard'],
        ['http://{tenant_id}.localhost/accounting',      'Accounting dashboard (if feature:accounting)'],
        ['http://{tenant_id}.localhost/api/user',        'Sanctum-protected test endpoint'],
    ],
    col_widths=[Cm(7), Cm(9)],
    font_size=9,
)

add_heading(doc, '8.3 Common Artisan Commands', level=2)
add_code_block(doc, """# Landlord / Central
php artisan tenants                                   # list tenants
php artisan tenants:create <tenant_id>                # create new tenant (admin-only path)

# Inside tenant context (after SetTenancy or via /super-admin UI):
php artisan migrate --path=database/migrations/tenant --force
php artisan db:seed --class=PermissionAndRoleSeeder --force
php artisan db:seed --class=DefaultChartOfAccountsSeeder --force

# Cache + ops
php artisan optimize:clear
php artisan route:list --path=pos
php artisan config:cache

# Tenancy package
php artisan tenancy:list
php artisan tenancy:migrate

# Reverb (if used)
php artisan reverb:start""")

add_heading(doc, '8.4 Default Plans & Feature Map', level=2)
add_table(doc,
    ['Feature Key', 'Starter', 'Pro', 'Enterprise', 'Custom'],
    [
        ['pos',                  '✓', '✓', '✓', 'opt'],
        ['purchase',             '✓', '✓', '✓', 'opt'],
        ['manufacturing',        '✗', '✓', '✓', 'opt'],
        ['multi_warehouse',      '✗', '✓ (max 5)', '✓', 'opt'],
        ['accounting',           '✗', '✓', '✓', 'opt'],
        ['accounting_advanced',  '✗', '✓', '✓', 'opt'],
        ['stock_count',          '✗', '✓', '✓', 'opt'],
        ['reports_advanced',     '✗', '✓', '✓', 'opt'],
    ],
    col_widths=[Cm(5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)],
    font_size=9,
)

add_heading(doc, '8.5 Key Files To Remember', level=2)
add_table(doc,
    ['Task', 'Open This File First'],
    [
        ['Add/modify a landlord route',           'routes/web.php'],
        ['Add/modify a tenant route',             'routes/tenant.php'],
        ['Tweak tenancy behavior',                'config/tenancy.php'],
        ['Adjust auth/permissions',               'app/Providers/AuthServiceProvider.php'],
        ['Adjust RBAC matrix',                    'database/seeders/PermissionAndRoleSeeder.php'],
        ['Add a new plan/feature',                'database/seeders/PlanSeeder.php + app/Models/Tenant::FEATURE_*'],
        ['Edit sales/purchase invoice logic',     'app/Services/InvoiceService.php (66 KB)'],
        ['Edit POS behavior',                     'app/Livewire/PosPanel.php + app/Http/Controllers/PosShiftController.php'],
        ['Edit accounting auto-post',             'app/Services/Accounting/PostingService.php (43 KB)'],
        ['Edit reports / financial logic',        'app/Services/ReportingService.php / app/Services/Accounting/FinancialReportService.php'],
        ['Edit UI (RTL, theme)',                  'resources/views/layouts/app.blade.php + layouts/theme-css.blade.php'],
        ['Edit landlord UI (dark glass)',         'resources/views/landlord/layout.blade.php'],
        ['Setup new tenant provisioning',         'app/Http/Controllers/Landlord/SuperAdminController.php + app/Providers/TenancyServiceProvider.php'],
        ['Edit COA template',                     'database/seeders/DefaultChartOfAccountsSeeder.php (12 KB)'],
        ['Edit dev env',                          '.env.example / .env'],
        ['Edit production env sample',            '.env.production.example'],
    ],
    col_widths=[Cm(5.5), Cm(10.5)],
    font_size=9,
)

add_heading(doc, '8.6 Troubleshooting Recipe', level=2)
add_table(doc,
    ['Symptom', 'Check'],
    [
        ['500 on /super-admin/*',          'Check APP_KEY, run php artisan optimize:clear'],
        ['500 on tenant route',            'Check tenant DB exists, migrations applied, feature flag enabled'],
        ['Tenant login fails',             'Verify tenant domain resolves, tenant DB has users + roles seeded'],
        ['POS panel blank',                'Check Livewire scripts loaded, debug vite build if used'],
        ['Invoice not auto-posting',       'Check /accounting/posting-failures; verify PostingService logs'],
        ['No style applied',               'Verify CDN fallback works; check resources/views/layouts/theme-css.blade.php'],
        ['Chart of Accounts empty',        'Re-run DefaultChartOfAccountsSeeder on tenant DB'],
        ['Plan upgrade missing features',  'Re-check plan features JSON in plans table; clear central caches'],
        ['Routes ordering bug',            'See commits 3f52e9d, 8732779 for /create vs {id} placement'],
        ['Multi-warehouse transfer fails', 'Check from/to warehouses are different + product has stock at source'],
    ],
    col_widths=[Cm(5), Cm(11)],
    font_size=9,
)

add_heading(doc, '8.7 Production Deployment Checklist', level=2)
deploy = [
    'Set APP_DEBUG=false in .env.',
    'Set APP_ENV=production.',
    'Generate APP_KEY with php artisan key:generate.',
    'php artisan config:cache; route:cache; view:cache.',
    'Run php artisan migrate --path=database/migrations (landlord) — Stancl will run tenant DBs on demand via CreateDatabase + MigrateDatabase.',
    'Configure CENTRAL_DOMAINS and TENANT_DOMAIN_SUFFIX in .env.',
    'Configure mail (SMTP), queue (redis/database), session (redis) — currently file drivers.',
    'Back up central DB regularly (spatie/laravel-backup configured).',
    'Ensure DB user has CREATE DATABASE privilege (needed for CreateDatabase job).',
    'Set up DNS wildcard *.yourdomain.com OR proxy that routes tenant hostnames correctly.',
    'Configure Reverb for WebSockets if any realtime features are enabled.',
    'Run security audit: APP_DEBUG, CORS, rate limits, password policies.',
]
for d in deploy:
    doc.add_paragraph(d, style='List Bullet')

# Closing note
add_page_break(doc)
add_hr(doc)
add_para(doc,
    'END OF DOCUMENTATION.  This file is the single source of truth for the Magzani ERP project. '
    'Always consult this document before exploring files, before answering user questions, and before '
    'writing new features. When in doubt, look at Part 4 (Modules), Part 5 (Routes), or Part 6 (Business Logic).',
    italic=True,
    align=1,  # WD_ALIGN_PARAGRAPH.CENTER
    color=RGBColor(0x6B, 0x72, 0x80),
)

# Save
doc.save(OUT_PATH)
print(f"Part 7 + Part 8 appended → {OUT_PATH}")
print(f"Final file size:")
import os
size_kb = os.path.getsize(OUT_PATH) / 1024
print(f"  {size_kb:.1f} KB")
