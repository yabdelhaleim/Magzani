"""
Final post-processing pass: Renumber ALL section headers to use their parent Part number.
Strategy:
  For each Heading 1 "Part N: title":
    Walk every paragraph until next Heading 1.
    For each Heading 2 in scope: assign numbers N.1, N.2, N.3, ...
    For each Heading 3 inside that Heading 2: assign N.K.1, N.K.2, ...
    Existing prefix is stripped and replaced.

Special skip rules:
  - For Part 6 (Modules), the H2 entries are "Module X: ..." — keep them as-is.
  - For Part 7 (Routes), the H2s use dot-notated like "5.2.1 Auth" — strip prefix and renumber 7.1, 7.2.
  - For Part 8 (Business Logic), the H2s use "6.1 Tenancy Internals" — renumber 8.1, 8.2.
  - For Part 9 (Recent Changes), renumber 9.1, 9.2, 9.3.
  - For Part 10 (Quick Reference), renumber 10.1, 10.2, ..., 10.7.
"""
import re
from docx import Document
import os

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
doc = Document(IN_PATH)

paras = doc.paragraphs

# Find each H1 boundary
h1_indices = [i for i, p in enumerate(paras) if p.style.name == 'Heading 1']
print(f"Found {len(h1_indices)} H1 paragraphs")

def set_para_text_keep_first_run_format(p, new_text):
    """Replace whole text with new_text, keeping first run formatting."""
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for run in p.runs[1:]:
        run.text = ''

# Track which parts we want to renumber and their target part number
def get_part_number(h1_text):
    m = re.match(r'Part (\d+):', h1_text)
    return int(m.group(1)) if m else None

# Process each part
total_renames = 0
for idx, h1_idx in enumerate(h1_indices):
    h1_para = paras[h1_idx]
    h1_text = h1_para.text.strip()
    part_num = get_part_number(h1_text)
    if part_num is None:
        continue
    # Determine scope
    scope_start = h1_idx + 1
    scope_end = h1_indices[idx + 1] if idx + 1 < len(h1_indices) else len(paras)

    # Skip Part 6 (Modules) — H2s are "Module X: ..." style
    if 'Modules (Deep Dive)' in h1_text:
        print(f"Skipping Part {part_num} (Modules) — using Module X: scheme")
        continue

    # Walk paragraphs in scope; renumber H2 sequentially
    h2_counter = 0
    h3_counter = 0
    skip_until_next_h1 = False

    for i in range(scope_start, scope_end):
        p = paras[i]
        if p.style.name == 'Heading 2':
            h2_counter += 1
            h3_counter = 0
            old_text = p.text.strip()
            # Strip any leading "N.M ..." or "N.M.K ..." prefix
            # Find first space → text after it
            m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', old_text)
            if m:
                title = m.group(2)
            else:
                title = old_text
            new_text = f'{part_num}.{h2_counter} {title}'
            set_para_text_keep_first_run_format(p, new_text)
            total_renames += 1
        elif p.style.name == 'Heading 3':
            h3_counter += 1
            old_text = p.text.strip()
            # Strip numeric prefix if present
            m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', old_text)
            if m:
                title = m.group(2)
            else:
                title = old_text
            if h2_counter > 0:
                new_text = f'{part_num}.{h2_counter}.{h3_counter} {title}'
            else:
                new_text = title
            set_para_text_keep_first_run_format(p, new_text)
            total_renames += 1

print(f"\n{total_renames} section headers renumbered")

# After renumbering inside "Module X" headers (Module 1, Module 2 etc.) we want them to live under Part 6.
# Inside each Module, the H3s are "Files", "Key Fields", "Routes", "Business Notes" — renumber as 6.X.1 ... 6.X.4
# (Currently they have NO numbers — they're just "Files", etc.) So they look fine as titles.
# But "tenants" / "domains" etc. inside Part 5 (DB Schema) — those are H3 with proper numbering already.

# Force H3 inside Part 6 to have no leading number (keep as plain titles) — they already do.

doc.save(IN_PATH)
print(f"\nSaved: {IN_PATH}")
print(f"Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB")

# Verify a sample
print("\n=== Sample of new headings ===")
count = 0
for p in paras:
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
        prefix = '    ' if p.style.name == 'Heading 3' else '  '
        print(f'{prefix}[{p.style.name}] {p.text}')
        count += 1
        if count > 60:
            break
