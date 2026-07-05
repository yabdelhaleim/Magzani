"""
Add the new 'Public Pricing Page' module documentation to the DOCX.
Appends a H1/H2 section at the very end of the document (after Part 10).
Also updates the TOC hint + adds a note about pricing subdomain config.
"""
import sys, os
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
doc = Document(IN_PATH)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0B, 0x11, 0x20)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x4B, 0x5E)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_table(doc, headers, rows, header_bg='1E3A8A', header_fg=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = header_fg
        run.font.name = 'Calibri'
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_borders(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
            set_cell_borders(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    return table

# ==========================================================
# ADD NEW PART 11 — Pricing Page (Public Kayyan SaaS)
# ==========================================================
add_page_break(doc)
add_heading(doc, 'Part 11: Public Pricing Page (Kayyan SaaS Marketing)', level=1)

add_para(doc,
    'A standalone, public marketing pricing page for Kayyan SaaS — separate from the '
    'tenant dashboard. Designed to be hosted on its own subdomain (e.g. pricing.kayyan.com) '
    'and shared via external links (website, social media, email).').italic = True

# 11.1 Overview
add_heading(doc, '11.1 Overview', level=2)
add_para(doc, 'Route: GET /pricing | Route name: pricing.public | Controller: App\\Http\\Controllers\\PricingController@index')
add_para(doc, 'Auth required: No. Tenant context: No (always uses central DB connection).')
add_para(doc, 'Total page size: ~78 KB HTML (uncompressed).', italic=True)

add_heading(doc, '11.2 File Inventory', level=2)
files_lines = [
    ('New file', 'app/Http/Controllers/PricingController.php', 'Fetches plans + features, builds WhatsApp messages, passes data to view'),
    ('New file', 'config/pricing.php', 'Env-driven config (WhatsApp, demo URL, brand name, SEO defaults)'),
    ('New file', 'resources/views/pricing/layout.blade.php', 'Standalone RTL layout with inlined SVG defs + critical CSS + SEO meta'),
    ('New file', 'resources/views/pricing/index.blade.php', 'Actual 9-section marketing page (Hero / Features / Pricing / Comparison / Add-ons / FAQ / Final CTA / Footer)'),
    ('New file', 'docs/PRICING_PAGE.md', 'Reference doc with deployment instructions'),
    ('Modified', 'routes/web.php', 'Added GET /pricing → PricingController@index'),
    ('Modified', '.env.example', 'Documented KAYYAN_* env vars + CENTRAL_DOMAINS examples'),
    ('Modified', '.env.production.example', 'Added pricing.kayyan.com to CENTRAL_DOMAINS + KAYYAN_* values'),
    ('Modified', 'resources/views/Dashboard/dashboard.blade.php', 'Added conditional Upgrade Banner (shows for starter/pro tenants)'),
]
for kind, path, desc in files_lines:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'[{kind}] {path}')
    run.bold = True
    run.font.size = Pt(10)
    sub = p.add_run(f' — {desc}')
    sub.font.size = Pt(10)

# 11.3 Routing Architecture
add_heading(doc, '11.3 Routing Architecture', level=2)
add_para(doc, 'The route lives in routes/web.php (NOT routes/tenant.php). This is critical:')
add_para(doc,
    'routes/web.php is wrapped by RouteServiceProvider with the "web" middleware group + '
    'the "central.domains" middleware (CentralDomainsOnly). Since "central.domains" allows '
    'requests ONLY from hosts in config(tenancy.central_domains), and Stancl tenancy '
    'bootstrappers are NEVER applied to web.php routes, the pricing page:').italic = True
for b in [
    'Reads from the CENTRAL DB (default DB_CONNECTION — same as landlord operations)',
    'Returns 404 on any tenant subdomain (acme.localhost, etc.)',
    'Is fully cacheable + SEO-indexable',
    'Does NOT trigger InitializeTenancyByDomain or any DatabaseTenancyBootstrapper',
]:
    doc.add_paragraph(b, style='List Bullet')

# 11.4 Subdomain Setup
add_heading(doc, '11.4 Subdomain Deployment (Production)', level=2)
add_para(doc, 'To deploy on a dedicated subdomain like pricing.kayyan.com:')
deploy_steps = [
    '1. DNS: point pricing.kayyan.com A/CNAME to the same server as the main dashboard.',
    '2. Web server (Nginx/Apache): serve the same public/ document root for both domains.',
    '3. .env: add pricing.kayyan.com to CENTRAL_DOMAINS (comma-separated):',
    '   CENTRAL_DOMAINS=superdashboard.remotelly1.site,pricing.kayyan.com',
    '4. SSL: ensure the subdomain has a valid HTTPS certificate.',
    '5. Test: curl -I https://pricing.kayyan.com/pricing should return 200.',
]
for s in deploy_steps:
    p = doc.add_paragraph(s, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# 11.5 Data Source
add_heading(doc, '11.5 Data Source', level=2)
add_para(doc, 'Plans are loaded from the CENTRAL plans table:')
code_p = doc.add_paragraph()
code_run = code_p.add_run("""Plan::where('is_active', true)
    ->whereIn('slug', PricingController::PUBLIC_PLAN_SLUGS)  // ['starter','pro','enterprise']
    ->orderByRaw(\"FIELD(slug, 'starter', 'pro', 'enterprise')\")
    ->with(['featuresList' => fn($q) => $q->where('is_enabled', true)])
    ->get();""")
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)
code_p.paragraph_format.left_indent = Cm(0.5)

add_para(doc, 'Legacy plans (basic, pos, manufacturing, pro-x, custom) are HIDDEN from the public page.', italic=True)

# 11.6 CTA Configuration
add_heading(doc, '11.6 CTA Configuration (Env-Override)', level=2)
add_table(doc,
    ['Env Var', 'Default', 'Purpose'],
    [
        ['KAYYAN_WHATSAPP_NUMBER', '+966500000000', 'WhatsApp number used by كل Contact Sales CTA'],
        ['KAYYAN_DEMO_URL', 'Google Calendar public pattern', 'احجز ديمو CTA target (override with Calendly / Cal.com)'],
        ['KAYYAN_SIGNUP_URL', 'env(APP_URL)/super-admin/tenants/create', 'Start Free Trial CTA target'],
        ['KAYYAN_BRAND_NAME', '"كيان SaaS"', 'Brand name in title, JSON-LD, social cards'],
        ['KAYYAN_SEO_TITLE', 'الأسعار | كيان SaaS ...', '<title> tag override'],
        ['KAYYAN_SEO_DESCRIPTION', 'اختر باقة كيان SaaS...', '<meta name="description"> override'],
        ['KAYYAN_SEO_KEYWORDS', 'نظام إدارة أعمال, نقاط بيع...', '<meta name="keywords"> override'],
    ],
    col_widths=[Cm(4), Cm(7), Cm(5)],
    font_size=9,
)

# 11.7 Page Sections
add_heading(doc, '11.7 Page Sections (Top → Bottom)', level=2)
sections_table = [
    ['1', 'Sticky Top Bar', 'Kayyan K-logo + nav links + Contact + Start Free Trial CTAs'],
    ['2', 'Hero', 'H1 headline + subtitle + dual CTAs + trust badges'],
    ['3', 'Trust Strip', '4 stats (+500 شركات, +2,400 مستخدمين, +1.2M فاتورة, 99.9% uptime)'],
    ['4', '8 Features Grid', 'POS / المخازن / التصنيع / المحاسبة / التقارير / RBAC / مخزون الخشب / الدعم'],
    ['5', 'Pricing Cards', '3 cards (Starter/Pro/Enterprise) — each with WhatsApp CTA + Demo CTA'],
    ['6', 'Comparison Table', '8-row feature matrix (desktop) + accordion view (mobile)'],
    ['7', 'Add-ons', '4 mini-cards: تدريب / دعم VIP / تخصيص / ترحيل بيانات'],
    ['8', 'FAQ', '6 collapsible <details> Q&As (zero JS)'],
    ['9', 'Final CTA + Footer', 'Email-less CTA + 4-column footer + copyright'],
]
add_table(doc,
    ['#', 'Section', 'Purpose'],
    sections_table,
    col_widths=[Cm(1), Cm(4.5), Cm(11)],
    font_size=9,
)

# 11.8 Upgrade Banner (in tenant dashboard)
add_heading(doc, '11.8 In-App Upgrade Banner', level=2)
add_para(doc,
    'A conditional Upgrade Banner was added to resources/views/Dashboard/dashboard.blade.php '
    'at the top of the dash-v3 div. It shows ONLY for tenants whose plan_id is "starter" or "pro".')
add_para(doc, 'Trigger condition: @if(in_array($currentPlanSlug, ["starter", "pro"], true))', italic=True)
add_para(doc, 'Target: opens the public pricing subdomain in a new tab via KAYYAN_SIGNUP_URL.')
add_para(doc,
    'Styling: gradient indigo→amber background with animated blob decorations, '
    'responsive (stacks vertically on mobile), uses inline @push("styles") CSS.', italic=True)

# 11.9 SEO & Performance
add_heading(doc, '11.9 SEO & Performance', level=2)
add_para(doc, 'Pre-loaded SEO meta in pricing/layout.blade.php:')
seo_items = [
    '<title> + <meta name="description"> + <meta name="keywords">',
    'Open Graph tags (og:type, og:title, og:description, og:url, og:locale ar_SA, og:image)',
    'Twitter Card tags (summary_large_image)',
    'JSON-LD ProductCollection schema (pushed from index view via @stack("json-ld"))',
    'Canonical URL (auto-derived from current request origin)',
    'hreflang-ready (og:locale:alternate en_US)',
]
for i in seo_items:
    doc.add_paragraph(i, style='List Bullet')

add_para(doc, 'For higher traffic, add caching in PricingController:')
code2 = doc.add_paragraph()
code2_run = code2.add_run("""Cache::remember('pricing_page_payload', 3600, function () use ($request) {
    return [...];  // heavy queries / API calls
});""")
code2_run.font.name = 'Consolas'
code2_run.font.size = Pt(9)

# 11.10 Custom Plan Slugs
add_heading(doc, '11.10 Custom Plan Slugs (Public)', level=2)
add_para(doc,
    'PricingController::PUBLIC_PLAN_SLUGS defines which plans appear on the public page. '
    'To show more plans, edit the constant in the controller.')
add_para(doc, 'Active plans in central DB (per PlanSeeder):', bold=True)
add_para(doc, '• starter ($99) — POS + Purchase')
add_para(doc, '• pro ($299) — adds Manufacturing + Multi-warehouse (5) + Accounting + Accounting Advanced + Reports Advanced')
add_para(doc, '• enterprise ($599) — all features unlimited')

# 11.11 Related Documentation
add_heading(doc, '11.11 Related Documentation', level=2)
add_para(doc, '• docs/PRICING_PAGE.md — quick reference card with deployment steps')
add_para(doc, '• Part 4 (Modules) — existing /plan/upgrade inside tenant context is the legacy equivalent')
add_para(doc, '• Part 7 (Routes) — Route added to landlord scope (routes/web.php)')
add_para(doc, '• Plan model — app/Models/Plan.php')
add_para(doc, '• PlanFeature model — app/Models/PlanFeature.php')

# Save
doc.save(IN_PATH)
print(f'Added Part 11 — Pricing Page to DOCX')
print(f'Saved: {IN_PATH}')
print(f'Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB')
