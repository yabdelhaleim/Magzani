"""
Magzani ERP Doc - Part 4A
Detailed Modules: Auth/RBAC, Landlord, Products, Customers, Suppliers,
                  Sales, Purchases, Returns
"""
import sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from _build_doc_part3 import (
    doc, add_heading, add_para, add_para_rtl, add_table, add_code_block,
    add_hr, add_page_break, OUT_PATH,
)
from docx.shared import Pt, Cm, RGBColor

def module_block(doc, name, name_ar, summary, files, key_fields, business_notes, url_prefix=None):
    """Render a single module section consistently."""
    add_heading(doc, name, level=2)
    add_para_rtl(doc, name_ar, bold=True, size=12, color=RGBColor(0x1E, 0x40, 0xAF))
    add_para(doc, summary)
    add_heading(doc, 'Files', level=3)
    for f in files:
        doc.add_paragraph(f, style='List Bullet')
    if key_fields:
        add_heading(doc, 'Key Fields / Models', level=3)
        add_table(doc, ['Entity', 'Notable Fields'], key_fields,
                  col_widths=[Cm(5), Cm(11)], font_size=9)
    if url_prefix is not None:
        add_heading(doc, 'Routes', level=3)
        add_para(doc, f'Base URL prefix: {url_prefix}  (see Part 5 for full details)')
    if business_notes:
        add_heading(doc, 'Business Notes', level=3)
        for n in business_notes:
            doc.add_paragraph(n, style='List Bullet')

# ============================================================
# PART 4 - MODULES
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 4: Modules (Deep Dive)', level=1)
add_para(doc,
    'Each module section below documents its Models, Controllers, Services, Livewire components, '
    'Views, main routes, and key business logic. Use the TOC to navigate directly to any module.'
)

# ----- MODULE 1: AUTH & RBAC -----
module_block(
    doc,
    'Module 1: Authentication & RBAC',
    'المصادقة والصلاحيات',
    'User authentication (login/logout) and per-tenant Role-Based Access Control with module+action permissions.',
    [
        'app/Http/Controllers/Auth/LoginController.php (throttled 5/min)',
        'app/Http/Controllers/Auth/RegisterController.php',
        'app/Http/Controllers/UserController.php (CRUD users, toggle-active)',
        'app/Http/Controllers/PermissionsController.php (roles, role-permissions, edit-user)',
        'app/Models/User.php (isAdmin, hasPermission, hasAnyPermission, hasRole)',
        'app/Models/Role.php / Permission.php',
        'app/Providers/AuthServiceProvider.php (Gate::before admin bypass, 19 accounting abilities)',
        'database/seeders/PermissionAndRoleSeeder.php (22KB comprehensive matrix)',
        'resources/views/auth/{login,register}.blade.php',
        'resources/views/users/{create,edit,index,show}.blade.php',
        'resources/views/permissions/{index,edit-user,print,roles}.blade.php',
    ],
    [
        ['User', 'id, name, email (unique), password, phone, is_active, role (admin|employee), remember_token, soft_deletes'],
        ['Role', 'id, name, display_name (AR), description, color, is_system'],
        ['Permission', 'id, name, display_name, description, module, action'],
    ],
    [
        'Defaults: roles "admin" + "employee" seeded per tenant. Admin gets ALL permissions via Gate::before in AuthServiceProvider.',
        'Login route: POST /login (rate-limited 5/min). Logout: POST /logout.',
        'Permission modules in seeder cover: sales.invoices.*, purchases.invoices.*, warehouse.products.*, warehouse.transfers.*, warehouse.movements.*, plus accounting group.',
        'User management routes are admin-only: /users/{create,edit,show,index,toggle-active}.',
        'Permissions UI: PermissionsController manages roles (create/edit/destroy + per-role permission editor) and per-user permission editor.',
        'Account route: /permissions/roles/{role}/permissions → update-role-permissions.',
    ],
    url_prefix='/login, /logout, /users/*, /permissions/*',
)

# ----- MODULE 2: MULTI-TENANCY / LANDLORD -----
module_block(
    doc,
    'Module 2: Multi-Tenancy / Landlord (Kayyan SaaS Admin)',
    'إدارة المستأجرين (لوحة المالك)',
    'Super-admin manages tenants and plans. Creates new tenants which triggers Stancl provisioning (DB creation + migrations + admin user seeding).',
    [
        'routes/web.php (38 lines — all landlord routes prefixed /super-admin)',
        'app/Http/Controllers/Landlord/SuperAdminController.php',
        'app/Models/Tenant.php (extends Stancl\\Tenancy\\Database\\Models\\Tenant, uses HasDatabase, HasDomains)',
        'app/Models/Plan.php / PlanFeature.php',
        'app/Providers/TenancyServiceProvider.php (events, middleware priority)',
        'config/tenancy.php (central tenancy config — 234 lines)',
        'resources/views/landlord/{layout,dashboard}.blade.php + plans/ + tenants/',
        'database/migrations/{create_tenants_table, create_domains_table, plans_table, plan_features_table, add_subscription_dates_to_tenants_table}.php',
        'database/seeders/PlanSeeder.php',
    ],
    [
        ['Tenant', 'id (UUID), data (JSON: plan_id, is_suspended), plan_expires_at, trial_ends_at (added 2026_06_05_000002)'],
        ['Plan', 'id, slug (unique), name, description, price decimal(10,2), billing_period, features (JSON), is_active'],
        ['PlanFeature', 'id, plan_id, feature_key, is_enabled, limit_value (nullable)'],
    ],
    [
        'Tenant creation flow: SuperAdminController::tenantsStore validates tenant_id (alpha_dash + lowercase + unique) → creates Tenant (fires TenantCreated event → queues CreateDatabase then MigrateDatabase jobs) → creates Domain {tenant_id}.{TENANT_DOMAIN_SUFFIX} → inside tenant()->run(...) seeds roles + creates admin user (admin@{tenant_id}.com / password) + attaches admin role.',
        'Tenant routes: GET/POST /super-admin/tenants, /create, /{id}/edit, /{id} (PUT), /{id}/toggle-status (POST), /{id} (DELETE).',
        'Plan CRUD: GET/POST /super-admin/plans, /create, /{id}/edit (PUT), /{id} (DELETE).',
        'Tenant model helper: publicUrl($path) builds tenant login URL shown on landlord dashboard.',
        'Central domain config: app.central_domains in config/tenancy.php (127.0.0.1, localhost, plus env CENTRAL_DOMAINS).',
    ],
    url_prefix='/super-admin/{dashboard, plans/*, tenants/*}',
)

# ----- MODULE 3: PRODUCTS & PRICING -----
module_block(
    doc,
    'Module 3: Products & Pricing',
    'المنتجات والتسعير',
    'Catalog of products with multi-unit pricing, supplier-specific prices, barcode generation, bulk price update, price history.',
    [
        'app/Http/Controllers/ProductController.php (18 KB — CRUD, search, low-stock)',
        'app/Http/Controllers/CategoryController.php',
        'app/Http/Controllers/PriceUpdateController.php (22 KB — bulk update UI)',
        'app/Http/Controllers/ProductPricingController.php',
        'app/Services/ProductService.php (50 KB)',
        'app/Services/ProductPricingService.php (15 KB)',
        'app/Services/PriceUpdateService.php (18 KB)',
        'app/Services/AdvancedPricingService.php (27 KB)',
        'app/Models/Product.php / Category.php / ProductBaseUnit.php / ProductBasePricing.php / ProductSellingUnit.php / ProductSalesUnit.php / ProductPriceHistory.php / PriceChangeHistory.php / UnitConversion.php / ProductWarehouse.php / SupplierProductPrice.php',
        'resources/views/products/{index,create,edit,show,barcode,bulk-price-update,price-history,units-statistics}.blade.php',
    ],
    [
        ['Product', 'name, code, sku, barcode, description, category_id, brand_id, unit_id, base_unit, base_unit_label, purchase_price, selling_price, min_selling_price, wholesale_price, tax_rate, default_discount, profit_margin, stock_alert_quantity, reorder_level, reorder_quantity, min_stock, max_stock, weight, dimensions, image, is_active, is_featured, has_expiry, track_serial, product_type, is_manufactured'],
    ],
    [
        'Product accessors: total_stock, total_available, total_reserved, stock_status (نفذ/منخفض/متوفر), is_low_stock, is_out_of_stock, base_selling_price, base_purchase_price, sellingUnitsWithPrices.',
        'Scopes: active, featured, search, byCategory, manufactured, lowStock, outOfStock.',
        'Routes (in routes/tenant.php lines 313-339): GET /products (index), GET /products/barcode/print, GET/POST /products/bulk-price-update (preview + apply), AJAX endpoints: categories-by-unit, by-unit-category, preview-smart-update, suggested-pricing, convert-unit-price, unit-details. CRUD: admin.only for create/store/edit/update/destroy. POST /products/{id}/update-price, GET /products/{id}/price-history.',
    ],
    url_prefix='/products/* (CRUD + bulk-price-update + barcode + 7 AJAX)',
)

# ----- MODULE 4: CUSTOMERS -----
module_block(
    doc,
    'Module 4: Customers',
    'العملاء',
    'Customer directory with opening/current balance, credit limit, statement view, and statement export.',
    [
        'app/Http/Controllers/CustomerController.php',
        'app/Services/CustomerService.php',
        'app/Models/Customer.php (very small fillable — balance updated via 2026_04_07 migration)',
        'app/Models/Payment.php (polymorphic for customer payments)',
        'app/Exports/CustomersExport.php + CustomerStatementExport.php',
        'resources/views/customers/{create,edit,index,show,statement}.blade.php',
    ],
    [
        ['Customer', 'id, name, phone, email, address, contact_person, opening_balance, current_balance (updated by migration 2026_04_07), credit_limit, is_active, code, soft_deletes'],
    ],
    [
        'Routes (lines 408-423): CRUD /customers, show /customers/{id}, statement /customers/{id}/statement, export /customers/{id}/statement/export (Excel).',
        'Customer has payments via polymorphic payments table (payable_type = SalesInvoice).',
        'IMPORTANT (per recent fixes): Use $transaction["type"] (array access), NOT $transaction->type — fixed in commit ec3b50c + ba001f6.',
    ],
    url_prefix='/customers/* (CRUD + statement + statement.export)',
)

# ----- MODULE 5: SUPPLIERS -----
module_block(
    doc,
    'Module 5: Suppliers',
    'الموردين',
    'Supplier directory with per-supplier product pricing, supplier payments, balance, statements, and statements exports.',
    [
        'app/Http/Controllers/SupplierController.php',
        'app/Http/Controllers/SupplierPaymentController.php',
        'app/Services/SupplierService.php',
        'app/Services/SupplierBalanceService.php',
        'app/Services/SupplierStatementService.php',
        'app/Models/Supplier.php / SupplierPayment.php / SupplierProductPrice.php',
        'app/Exports/SuppliersExport.php + SupplierStatementExport.php',
        'resources/views/suppliers/{create,edit,index,show,statement}.blade.php',
    ],
    [
        ['Supplier', 'id, name, phone, email, address, contact_person, opening_balance, current_balance, is_active, code, soft_deletes'],
        ['SupplierPayment', 'id, supplier_id, invoice_id, amount, payment_date, payment_method, notes, journal_entry_id, created_by'],
        ['SupplierProductPrice', 'id, supplier_id, product_id, price, min_quantity, valid_from, valid_to, notes'],
    ],
    [
        'Routes (lines 430-445): CRUD /suppliers, show /suppliers/{id}, statement /suppliers/{id}/statement, export /suppliers/{id}/statement/export.',
        'SupplierController resolved via InvoiceService (66 KB) for invoice-side supplier balance updates.',
        'Null-safe Company access added in statement view (commit ba001f6).',
    ],
    url_prefix='/suppliers/* (CRUD + statement + statement.export)',
)

# ----- MODULE 6: SALES -----
module_block(
    doc,
    'Module 6: Sales Invoices & Returns',
    'فواتير المبيعات والمرتجعات',
    'Sales invoicing with multi-tax, discounts, payment terms, installment tracking; sales returns tied to original invoice.',
    [
        'app/Http/Controllers/SalesController.php (19 KB — main sales controller)',
        'app/Http/Controllers/SalesReturnsController.php (5 KB)',
        'app/Services/InvoiceService.php (66 KB — core sales/purchase invoice engine + confirm/post/etc.)',
        'app/Services/ReturnService.php (21 KB)',
        'app/Models/SalesInvoice.php / SalesInvoiceItem.php / SalesReturn.php / SalesReturnItem.php',
        'resources/views/invoices/sales/{create,edit,index,show,print,store,Print,thermal_receipt}.blade.php',
        'resources/views/invoices/sales-returns/{create,index,show}.blade.php',
    ],
    [
        ['SalesInvoice', 'invoice_number, customer_id, warehouse_id, invoice_date, due_date, subtotal, discount_type/value/amount, tax_rate/amount, shipping_cost, other_charges, total, paid, status (draft/confirmed/cancelled), payment_status (paid/partial/unpaid), confirmed_by/at, shift_id, source (pos/manual), payment_method, journal_entry_id, cogs_entry_id'],
        ['SalesInvoiceItem', 'invoice_id, product_id, quantity, unit, base_quantity, unit_price, discount_*, tax_*, total, cost_price, profit'],
        ['SalesReturn', 'return_number, invoice_id, customer_id, warehouse_id, return_date, subtotal, total, notes, reason, status, shift_id'],
    ],
    [
        'Accessors: remaining (total - paid), payment_percentage, is_fully_paid, is_partially_paid.',
        'Scopes: paid, unpaid, partiallyPaid, confirmed.',
        'Routes (feature:pos gated, lines 107-120): CRUD /invoices/sales (admin-only writes), /invoices/sales/{id}/print.',
        'Confirmation flow (inside InvoiceService::confirmInvoice — DB transaction + lockForUpdate): decrement inventory → record InventoryMovement type=sale → update Customer.current_balance → auto-post JournalEntry (Sales) via PostingService → set status=confirmed, confirmed_by/at → generate cogs_entry_id.',
        'Thermal receipt view exists: invoices/sales/thermal_receipt.blade.php.',
    ],
    url_prefix='/invoices/sales/* (CRUD + print), /invoices/sales-returns/*',
)

# ----- MODULE 7: PURCHASES -----
module_block(
    doc,
    'Module 7: Purchase Invoices & Returns',
    'فواتير المشتريات والمرتجعات',
    'Purchase invoicing, supplier balance updates, purchase returns; gated behind feature:purchase plan flag.',
    [
        'app/Http/Controllers/PurchaseInvoiceController.php (5.6 KB)',
        'app/Http/Controllers/PurchaseReturnController.php (6.7 KB)',
        'app/Services/PurchaseInvoiceService.php (11 KB)',
        'app/Services/PurchaseReturnService.php (12 KB)',
        'app/Models/PurchaseInvoice.php / PurchaseInvoiceItem.php / PurchaseReturn.php / PurchaseReturnItem.php',
        'app/Exports/PurchaseInvoicesExport.php + PurchaseInvoiceDetailsExport.php',
        'resources/views/invoices/purchases/{create,edit,index,show,print}.blade.php',
        'resources/views/invoices/purchase-returns/index.blade.php',
    ],
    [
        ['PurchaseInvoice', 'invoice_number, supplier_id, warehouse_id, invoice_date, due_date, subtotal, discount_*, tax_*, total, paid, status, payment_status, confirmed_by/at, journal_entry_id'],
    ],
    [
        'Routes (feature:purchase + auth, lines 360-372): CRUD /invoices/purchases/*, POST /invoices/purchases/{id}/confirm, GET /print, GET /export (list), GET /{id}/exportSingle.',
        'Purchase-returns routes (lines 392-401): CRUD + AJAX /invoices/purchase-returns/available-items.',
        'Confirmation: increase inventory + record InventoryMovement type=purchase + update Supplier.current_balance + auto-post JournalEntry (Purchase) + mark confirmed.',
    ],
    url_prefix='/invoices/purchases/* (feature:purchase), /invoices/purchase-returns/*',
)

doc.save(OUT_PATH)
print(f"Part 4A appended → {OUT_PATH}")
