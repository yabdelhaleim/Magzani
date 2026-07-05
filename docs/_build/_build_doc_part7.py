"""
Magzani ERP Doc - Part 6
Business Logic Deep-Dive: POS Lifecycle, Accounting Posting Flow, RBAC, Plans
"""
import sys
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from _build_doc_part6 import (
    doc, add_heading, add_para, add_para_rtl, add_table, add_code_block,
    add_hr, add_page_break, OUT_PATH,
)
from docx.shared import Pt, Cm, RGBColor

# ============================================================
# PART 6 - BUSINESS LOGIC
# ============================================================
add_page_break(doc)
add_heading(doc, 'Part 6: Business Logic Deep-Dive', level=1)

# ----- Tenancy Internals -----
add_heading(doc, '6.1 Tenancy Internals', level=2)

add_heading(doc, '6.1.1 Tenant Model Deep-Dive', level=3)
add_para(doc, 'app/Models/Tenant.php extends Stancl\\Tenancy\\Database\\Models\\Tenant. It implements TenantWithDatabase, uses HasDatabase, HasDomains traits, and adds:')
add_code_block(doc, """// app/Models/Tenant.php (key features)

// Constants for plan slugs
const PLAN_STARTER    = 'starter';
const PLAN_PRO        = 'pro';
const PLAN_ENTERPRISE = 'enterprise';
const PLAN_CUSTOM     = 'custom';

// Feature keys
const FEATURE_POS                 = 'pos';
const FEATURE_MANUFACTURING       = 'manufacturing';
const FEATURE_MULTI_WAREHOUSE     = 'multi_warehouse';
const FEATURE_ACCOUNTING          = 'accounting';
const FEATURE_ACCOUNTING_ADVANCED = 'accounting_advanced';
const FEATURE_STOCK_COUNT         = 'stock_count';
const FEATURE_PURCHASE            = 'purchase';
const FEATURE_REPORTS_ADVANCED    = 'reports_advanced';

// Dynamic: get plan via data.plan_id (JSON)
public function getPlanAttribute() {
    // returns Plan from central DB or synthetic custom plan
}

// Feature check
public function hasFeature(string $key): bool {
    // checks plan()->features[$key] or custom_features
}

// Subscription status
public function isOnTrial(): bool;
public function isActive(): bool;
public function hasExpired(): bool;

// URL helper
public function publicUrl(string $path = '/'): string {
    // returns https://{primaryDomain}{$path}
}""")

add_heading(doc, '6.1.2 TenancyServiceProvider', level=3)
add_para(doc, 'Registered in config/app.php line ~171. Does three things:')
for it in [
    '1. Bootstraps tenant lifecycle events (TenantCreated → CreateDatabase + MigrateDatabase jobs; TenantDeleted → DeleteDatabase).',
    '2. Sets tenancy middleware to highest priority (preventing route middleware conflicts) via makeTenancyMiddlewareHighestPriority().',
    '3. Marks tenancy route loading as the responsibility of RouteServiceProvider (commented out in TenancyServiceProvider to avoid UrlGenerator state desync).',
]:
    doc.add_paragraph(it, style='List Bullet')

add_heading(doc, '6.1.3 Feature Middleware', level=3)
add_para(doc, 'Used as feature:pos, feature:manufacturing, feature:multi_warehouse, feature:accounting, feature:accounting_advanced, feature:purchase, feature:warehouses, feature:stock_count.')
add_code_block(doc, """// Pseudo-code of feature middleware (HTTP middleware)
public function handle($request, Closure $next, string $feature) {
    if (! tenant()->hasFeature($feature)) {
        abort(403, 'Feature not available in your plan');
    }
    return $next($request);
}""")

add_heading(doc, '6.1.4 Tenant Provisioning (Recap)', level=3)
add_para(doc, 'When SuperAdmin creates a tenant (routes/web.php → SuperAdminController@tenantsStore):')
flow = [
    'Validate tenant_id (alpha_dash, lowercase, unique).',
    'Tenant::create([id => $tenant_id]) → fires TenantCreated event.',
    'Stancl dispatches CreateDatabase job (creates MySQL DB named tenant{tenant_id}).',
    'On DB create success, dispatches MigrateDatabase job.',
    'MigrateDatabase runs php artisan migrate --path=database/migrations/tenant --force --realpath.',
    'Domain::create([domain => "{tenant_id}.{TENANT_DOMAIN_SUFFIX}", tenant_id => $tenant_id]).',
    '$tenant->run(...) closure runs inside tenant context:',
    '   • PermissionAndRoleSeeder → admin + employee roles + all permissions.',
    '   • User::create([email => "admin@{tenant_id}.com", password => bcrypt("password"), role => "admin", is_active => true]).',
    '   • $user->attachRole(admin).',
    'Return redirect to /super-admin/tenants with success flash.',
]
for s in flow:
    doc.add_paragraph(s, style='List Bullet')

# ----- POS -----
add_heading(doc, '6.2 POS Shift Lifecycle', level=2)

add_heading(doc, '6.2.1 Status State Machine', level=3)
add_code_block(doc, """STATUS_OPEN       = 'open'      // shift in progress, can record sales
STATUS_CLOSED     = 'closed'    // shift closed normally by cashier
STATUS_AUTO_CLOSED = 'auto_closed' // system-closed: stale from previous day or replaced""")

add_heading(doc, '6.2.2 Open Shift Flow', level=3)
add_code_block(doc, """// PosShiftController::open(Request $request)
// 1. validate opening_balance (numeric, min:0)
// 2. ensure no active shift: PosShift::getActiveShift($request->user()->id) === null
// 3. $previousAutoClosed = PosShift::autoCloseStaleShays()  // closes stale shifts from past days
// 4. $shift = PosShift::create([
//      user_id => auth()->id(),
//      opened_at => now(),
//      opening_balance => $request->opening_balance,
//      status => STATUS_OPEN,
//      total_sales => 0, total_returns => 0, sales_count => 0, returns_count => 0,
//    ])
// 5. return redirect('/pos')  // Livewire reads active shift""")

add_heading(doc, '6.2.3 During Sale (Livewire)', level=3)
add_para(doc, 'PosPanel Livewire component renders the POS screen. On sale confirmation:')
add_code_block(doc, """// PosPanel::confirmSale()
// 1. InvoiceService::confirmInvoice($invoice) inside DB::transaction
//    a. lockForUpdate on product_warehouse rows (prevent race)
//    b. decrement product_warehouse.quantity, available_quantity
//    c. record InventoryMovement{type='sale'}
//    d. update Customer.current_balance (if credit sale) or mark paid
//    e. PostingService::postSalesInvoice($invoice)
//       → creates JournalEntry{source=system, status=posted}
//       → creates 2+ JournalEntryLines (Dr Cash/AR, Cr Sales, possibly COGS Dr / Cr Inventory)
//       → sets $invoice->journal_entry_id
//       → sets $invoice->cogs_entry_id (Cost of Goods Sold entry)
//    f. set $invoice->status=confirmed, confirmed_by, confirmed_at
// 2. PosShift::recalculateTotals() sums all confirmed POS sales for shift
//    - total_sales, sales_count, net_sales (total - returns)
//
// Livewire emits 'saleCompleted' event → UI updates cart/grid.""")
add_para(doc, 'Failed postings (e.g. invalid COGS) are logged to accounting_posting_failures and can be retried via /accounting/posting-failures.')

add_heading(doc, '6.2.4 Close Shift Flow', level=3)
add_code_block(doc, """// PosShiftController::close(Request $request)
// 1. load active shift
// 2. show close form (closeView) — shift summary + cash expected
// 3. on POST close:
//    a. validate closing_balance_actual (numeric)
//    b. $shift->calculateExpectedBalance() → opening_balance + SUM(cash sales)
//    c. $shift->update([
//         closing_balance_actual => $actual,
//         closing_balance_expected => $expected,
//         status => STATUS_CLOSED,
//         closed_at => now(),
//       ])
//    d. $shift->computeAndSaveDifference() → cash_difference = actual - expected
// 4. redirect to /pos/history (showing all closed shifts including this one)""")

add_heading(doc, '6.2.5 X-Report & Z-Report', level=3)
for s in [
    'X-Report (pos.xreport): live snapshot of current shift totals without closing it.',
    'Z-Report (pos.shift.zreport at /pos/shift/{id}/z-report): final end-of-day report for a closed shift — same data as X but read-only.',
    'History (pos.history): all past shifts paginated.',
]:
    doc.add_paragraph(s, style='List Bullet')

# ----- Accounting Flow -----
add_page_break(doc)
add_heading(doc, '6.3 Accounting Posting Flow', level=2)

add_heading(doc, '6.3.1 PostingService (43 KB) - The Auto-Posting Engine', level=3)
add_para(doc, 'PostingService lives at app/Services/Accounting/PostingService.php. It is the central entry point that turns business transactions into double-entry JournalEntries. Every confirmed invoice / payment / manufacturing completion flows through here.')
add_para(doc, 'PostingService methods (non-exhaustive):')
add_table(doc,
    ['Method', 'Source Transaction', 'Entries Created'],
    [
        ['postSalesInvoice($invoice)',     'Confirmed sales invoice',           'Dr Cash/AR, Cr Sales (+ optional COGS Dr Inventory / Cr Inventory)'],
        ['postPurchaseInvoice($invoice)',  'Confirmed purchase invoice',        'Dr Inventory, Cr Cash/AP (+ optional VAT)'],
        ['postPayment($payment)',          'Customer or supplier payment',      'Dr/Cash + Cr AR for receipts; Dr AP + Cr Cash for payments'],
        ['postRefund($refund)',            'Customer refund',                   'Dr Sales Returns, Cr Cash'],
        ['postManufacturingCompletion($mo)', 'Manufacturing order completed',   'Dr Inventory (FG), Cr Inventory (Components) + variance'],
        ['postExpense($expense)',          'Expense record',                    'Dr Expense, Cr Cash'],
        ['postCashTransaction($txn)',      'Cash deposit or withdrawal',        'Dr/Cash, Cr/Bank or vice versa'],
        ['postFixedAssetDepreciation($fa)', 'Depreciation entry',               'Dr Depreciation Expense, Cr Accumulated Depreciation'],
        ['postVatSettlement($period)',     'Periodic VAT settlement',           'Dr VAT Payable, Cr VAT Receivable (or reverse)'],
    ],
    col_widths=[Cm(5), Cm(5), Cm(6)],
    font_size=9,
)

add_heading(doc, '6.3.2 Posting Failure Handling', level=3)
add_para(doc, 'Every post*() method wraps the journal insert in a try/catch. On failure:')
flow_fail = [
    '1. Catch Throwable.',
    '2. Insert AccountingPostingFailure row:',
    '   {',
    '     transaction_type: <class>,      // e.g. SalesInvoice',
    '     reference_id: <id>,',
    '     error_message: <message>,',
    '     retry_count: 0,',
    '     resolved_at: null',
    '   }',
    '3. Original transaction still goes through (status confirmed), but linked journal_entry_id stays null until retry succeeds.',
    '4. /accounting/posting-failures shows pending failures with retry/resolve actions.',
    '5. Failed invoices are still reversible (their stock movement happened); only the GL posting retries.',
]
for s in flow_fail:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '6.3.3 Year-End Closing Wizard', level=3)
add_para(doc, '/accounting/fiscal/years/{year}/year-end-close triggers YearEndClosingController which:')
add_code_block(doc, """// Pseudo-code
public function yearEndClose(FiscalYear $year) {
    DB::transaction(function () use ($year) {
        // 1. Verify all periods in year are closed
        // 2. Post net income/loss to retained earnings
        $retainedEarnings = Account::where('code', 'retained_earnings')->first();
        JournalEntry::create([
            'source' => 'system',
            'description' => 'Year-end closing for ' . $year->name,
            'lines' => [...],
        ]);
        // 3. Mark year as closed
        $year->update(['is_closed' => true, 'closed_at' => now()]);
    });
}""")

# ----- RBAC -----
add_page_break(doc)
add_heading(doc, '6.4 Role-Based Access Control (RBAC)', level=2)

add_heading(doc, '6.4.1 User Model Key Methods', level=3)
add_code_block(doc, """// app/Models/User.php
public function isAdmin(): bool {
    return $this->role === 'admin';          // legacy shortcut
        || $this->hasAllPermissionsViaRole(); // RBAC-based check
}

public function hasPermission(string $perm): bool {
    // returns true if any role has this permission
    // super-bypass: Gate::before returns true if isAdmin()
}

public function hasAnyPermission(array $perms): bool { ... }
public function hasRole(string $name): bool { ... }
public function allPermissions(): \Illuminate\Support\Collection { ... }
public function roles(): BelongsToMany { return $this->belongsToMany(Role::class); }
public function permissions(): BelongsToMany { return $this->belongsToMany(Permission::class); }""")

add_heading(doc, '6.4.2 Gate::before in AuthServiceProvider', level=3)
add_code_block(doc, """// app/Providers/AuthServiceProvider.php
public function boot() {
    $this->registerPolicies();
    
    Gate::before(function ($user, $ability) {
        if ($user->isAdmin()) {
            return true;  // admins bypass all checks
        }
    });
    
    // warehouse transfers
    Gate::define('warehouse.transfers.read', fn($u) => $u->hasPermission('warehouse.transfers.read'));
    Gate::define('warehouse.transfers.create', fn($u) => $u->hasPermission('warehouse.transfers.create'));
    Gate::define('warehouse.transfers.update', fn($u) => $u->hasPermission('warehouse.transfers.update'));
    Gate::define('warehouse.transfers.delete', fn($u) => $u->hasPermission('warehouse.transfers.delete'));
    
    // users
    Gate::define('users.permissions', fn($u) => $u->hasPermission('users.permissions'));
    
    // accounting (19 abilities defined dynamically)
    foreach ([
        'accounting.dashboard', 'accounting.coa.read', 'accounting.coa.write',
        'accounting.journal.read', 'accounting.journal.write', 'accounting.journal.post', 'accounting.journal.reverse',
        'accounting.vouchers.read', 'accounting.vouchers.write',
        'accounting.reports.read', 'accounting.fiscal.read', 'accounting.fiscal.write',
        'accounting.settings.read', 'accounting.settings.write',
        'accounting.posting-failures.read', 'accounting.posting-failures.write',
        // ... 19 in total
    ] as $ability) {
        Gate::define($ability, function ($u) use ($ability) {
            return $u->hasPermission($ability);
        });
    }
}""")

add_heading(doc, '6.4.3 Permission Seeding (22 KB matrix)', level=3)
add_para(doc, 'PermissionAndRoleSeeder creates a comprehensive matrix. Key permission groups:')
perm_groups = [
    'sales.invoices.{create,read,update,delete,print}',
    'sales.returns.{create,read,delete}',
    'purchases.invoices.{create,read,update,delete}',
    'purchases.returns.{...}',
    'warehouse.products.{create,read,update,delete}',
    'warehouse.transfers.{create,read,update,delete}',
    'warehouse.movements.{read}',
    'accounting.* (19 abilities — see AuthServiceProvider)',
    'users.permissions',
]
for p in perm_groups:
    doc.add_paragraph(p, style='List Bullet')

# ----- Plans -----
add_page_break(doc)
add_heading(doc, '6.5 Subscriptions & Plans', level=2)

add_heading(doc, '6.5.1 Default Plans (PlanSeeder)', level=3)
add_table(doc,
    ['Slug', 'Price', 'Features (sample)'],
    [
        ['starter',    '$99/month',  'pos, purchase'],
        ['pro',        '$299/month', 'pos, purchase, manufacturing, multi_warehouse (limit 5), accounting, accounting_advanced, reports_advanced'],
        ['enterprise', '$599/month', 'all features unlimited'],
        ['basic',      '$19 (legacy)', 'base only'],
        ['pos',        '$39 (legacy)', 'pos'],
        ['manufacturing','$79 (legacy)','manufacturing'],
        ['pro-x',      '$500 (legacy)','pro extras'],
        ['custom',     'price=0',     'special: data.custom_features array controls'],
    ],
    col_widths=[Cm(3), Cm(3.5), Cm(9.5)],
    font_size=9,
)

add_heading(doc, '6.5.2 Plan Resolution Flow', level=3)
add_code_block(doc, """// app/Models/Tenant.php - getPlanAttribute()
public function getPlanAttribute(): Plan|CustomPlan {
    $planId = $this->data['plan_id'] ?? null;
    if ($planId === 'custom' || $planId === null) {
        // fallback: return Plan::custom with custom_features
        return $this->resolveCustomPlan();
    }
    return Plan::where('slug', $planId)->first() ?? $this->resolveCustomPlan();
}

public function hasFeature(string $key): bool {
    $plan = $this->plan;
    if (! $plan) return false;
    if ($plan->slug === 'custom') {
        return in_array($key, $this->data['custom_features'] ?? []);
    }
    return (bool) ($plan->features[$key] ?? false);
}""")

add_heading(doc, '6.5.3 Plan Limits', level=3)
add_para(doc, 'Some features carry numeric limits via PlanFeature.limit_value (e.g. multi_warehouse = 5 for pro plan). The code that checks these is typically in the relevant service / controller, not the middleware.')

add_heading(doc, '6.5.4 Subscription Status', level=3)
add_para(doc, 'Tenant table columns added by migration 2026_06_05_000002:')
add_table(doc,
    ['Column', 'Purpose'],
    [
        ['plan_expires_at', 'When current plan expires — used by isActive()'],
        ['trial_ends_at',   'End of trial period — used by isOnTrial()'],
    ],
    col_widths=[Cm(5), Cm(11)],
)
add_para(doc, 'Per Tenant::isActive(): returns plan_expires_at === null OR plan_expires_at > now(). Per Tenant::hasExpired(): opposite.')

add_heading(doc, '6.5.5 Plan Upgrade UI', level=3)
add_para(doc, 'GET /plan/upgrade (PlanController) lets a tenant user see available plans and trigger an upgrade request. Final payment integration is not implemented in this codebase (UI only — extends to a real payment gateway as a follow-up).')

# ----- Reports & Exports -----
add_page_break(doc)
add_heading(doc, '6.6 Reports & Exports', level=2)

add_heading(doc, '6.6.1 Excel Exports (Maatwebsite)', level=3)
add_table(doc,
    ['Class', 'Source Report'],
    [
        ['CustomersExport',         '/customers (list)'],
        ['SuppliersExport',         '/suppliers (list)'],
        ['InventoryReportExport',   '/reports/inventory/export'],
        ['ProfitLossReportExport',  '/reports/profit-loss/export'],
        ['FinancialReportExport',   '/reports/financial/export'],
        ['CustomerStatementExport', '/customers/{id}/statement/export'],
        ['SupplierStatementExport', '/suppliers/{id}/statement/export'],
        ['PurchaseInvoicesExport',  '/invoices/purchases/export'],
        ['PurchaseInvoiceDetailsExport', '/invoices/purchases/{id}/exportSingle'],
    ],
    col_widths=[Cm(6), Cm(10)],
    font_size=9,
)

add_heading(doc, '6.6.2 PDF Generation (DomPDF)', level=3)
add_para(doc, 'Used in: invoices/sales/{print,thermal_receipt}, invoices/purchases/print, customers/statement, suppliers/statement, accounting/vouchers/{receipt,payment}/print, stock-counts/{id}/print, warehouse-orders/{inbound,outbound}/{id}/print, permissions/print.')

add_heading(doc, '6.6.3 Reporting Service Hierarchy', level=3)
add_para(doc, 'reportingService (app/Services/ReportingService.php, 23KB) handles operational reports. financialReportService (in Accounting/) handles accounting reports with sub-methods for:')
rf = [
    'trialBalance($date)',
    'incomeStatement($from, $to)',
    'comparativeIncome($currentFrom, $currentTo, $compareFrom, $compareTo)',
    'balanceSheet($date)',
    'generalLedger($accountId, $from, $to)',
    'partnerLedger($partnerType, $partnerId, $from, $to)',
    'auditTrail($from, $to)',
    'aging($type = AR|AP, $date)',
    'financialRatios($from, $to)',
    'vatSettlement($period)',
]
for m in rf:
    doc.add_paragraph(m, style='List Bullet')

# ----- AppServiceProvider -----
add_heading(doc, '6.7 AppServiceProvider (Cross-Cutting)', level=2)
add_code_block(doc, """// app/Providers/AppServiceProvider.php
public function boot() {
    // 1. Tenant observer
    Tenant::observe(TenantObserver::class);
    
    // 2. View composer — inject plan features into EVERY view
    try {
        $planFeatures = tenant()?->plan?->features ?? collect();
        View::composer('*', function ($view) use ($planFeatures) {
            $view->with('planFeatures', $planFeatures);
        });
    } catch (\\Throwable $e) {
        // central-domain context (no tenant) — silently skip
    }
    
    // 3. Register any global scopes / custom macros here
}""")

add_heading(doc, '6.8 Custom Middleware Summary', level=2)
add_table(doc,
    ['Middleware', 'Purpose'],
    [
        ['auth', 'Standard Laravel session auth'],
        ['guest', 'Standard Laravel guest'],
        ['throttle:5,1', 'Rate-limit 5 attempts per minute (login)'],
        ['throttle:30,1', 'Rate-limit 30 per minute (journal entries)'],
        ['admin.only', 'Restrict route to admin role'],
        ['role', 'Custom role-based gate'],
        ['feature:<key>', 'Check tenant hasFeature(<key>)'],
        ['PreventAccessFromCentralDomains', 'Stancl — tenant routes block central hostnames'],
        ['InitializeTenancyByDomain', 'Stancl — resolve tenant from host'],
    ],
    col_widths=[Cm(5), Cm(11)],
    font_size=9,
)

doc.save(OUT_PATH)
print(f"Part 6 (Business Logic) appended → {OUT_PATH}")
