"""
Append a v2 addendum to Part 11 in the DOCX documenting:
1. New CMS editor at /super-admin/pricing-settings
2. Removal of all WhatsApp CTAs from the public page
3. The pricing_settings CMS table + pricing_settings.json structure

Strategy: append a "11.12 What's New in v2 (CMS editor + WhatsApp removed)"
section to Part 11.
"""
import sys, os
sys.path.insert(0, r"C:\MAGZANIV6\Magzani\docs")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

IN_PATH = r"C:\MAGZANIV6\Magzani\docs\Magzani_ERP_Full_Documentation.docx"
doc = Document(IN_PATH)

# ---------- helper imports from existing scripts ----------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0B, 0x11, 0x20)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x4B, 0x5E)
    return h

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_table(doc, headers, rows, header_bg='1E3A8A', header_fg=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths=None, font_size=10):
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = header_fg
        run.font.name = 'Calibri'
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_borders(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
            set_cell_borders(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    return table

# Append to the very end of the document
add_heading(doc, '11.12 v2 Update — CMS Editor + WhatsApp Removal (2026-07-04)', level=2)

add_heading(doc, '11.12.1 What Changed', level=3)
for s in [
    'Added a full CMS admin form at /super-admin/pricing-settings to edit the marketing copy around pricing cards (hero, trust stats, features, addons, FAQs, footer, etc.) — no code changes needed to tweak the public pricing page.',
    'Removed ALL WhatsApp CTAs from the public page (top bar, hero, plan cards, final CTA, footer social). The only remaining CTA per card is now a single "احجز ديمو مجاني" → demo URL.',
    'Added new sidebar link in the landlord dashboard: "إعدادات صفحة الأسعار" with a CMS badge, between Plans and Tenants.',
    'Removed unused .k-btn-whatsapp CSS class from pricing/layout.blade.php.',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '11.12.2 New CMS Architecture', level=3)
add_heading(doc, '11.12.2.1 Storage Model', level=3)
doc.add_paragraph(
    'A single-row key/value table in the central DB (pricing_settings, key = "page") stores the entire content tree as JSON. '
    'A model-level cache (60 minutes) prevents repeated deserialization; write operations bust the cache automatically.',
    style='List Bullet'
)
doc.add_paragraph(
    'The model exposes PricingSetting::current() (with fallback to defaults()) and PricingSetting::saveAll($array).',
    style='List Bullet'
)

add_heading(doc, '11.12.2.2 Route & Controller', level=3)
add_table(doc,
    ['URL', 'Name', 'Action'],
    [
        ['GET  /super-admin/pricing-settings', 'super-admin.pricing-settings.index', 'PricingSettingsController@index — render form'],
        ['POST /super-admin/pricing-settings', 'super-admin.pricing-settings.update', 'PricingSettingsController@update — sanitize + save'],
    ],
    col_widths=[Cm(6), Cm(6), Cm(4)],
    font_size=9,
)

add_heading(doc, '11.12.2.3 Editable Sections', level=3)
add_table(doc,
    ['#', 'Section', 'Editable fields'],
    [
        ['1', 'Hero', 'eyebrow, headline (HTML), subtitle, primary+secondary CTA label+URL, trust badges, visibility'],
        ['2', 'Trust stats', '4 cards (icon + value + label), visibility'],
        ['3', 'Features grid', 'eyebrow, heading, subtitle, 8 cards, visibility'],
        ['4', 'Add-ons', 'eyebrow, heading, 4 mini-cards, visibility'],
        ['5', 'FAQ', 'eyebrow, heading, 6 Q&A pairs, visibility'],
        ['6', 'Final CTA', 'eyebrow, heading, subtitle, button label, visibility'],
        ['7', 'General', 'Hard on/off kill switch + per-card CTA label'],
    ],
    col_widths=[Cm(1), Cm(3), Cm(12)],
    font_size=9,
)

add_heading(doc, '11.12.2.4 Data Schema', level=3)
code = doc.add_paragraph()
code_run = code.add_run("""pricing_settings (
  id BIGINT PRIMARY KEY,
  `key` VARCHAR(100) UNIQUE,        -- always "page" for v1
  value JSON,                       -- e.g. {"hero":{"headline":"..."},"features_grid":{"items":[...]},...}
  created_at, updated_at
)""")
code_run.font.name = 'Consolas'
code_run.font.size = Pt(9)
code.paragraph_format.left_indent = Cm(0.5)

add_heading(doc, '11.12.3 WhatsApp Removal — Diff', level=3)
add_table(doc,
    ['Location', 'Before', 'After'],
    [
        ['Top bar', 'WhatsApp icon button → wa.me/...', 'Removed entirely'],
        ['Hero', 'Secondary CTA "تحدث مع المبيعات" → wa.me', 'Replaced by "احجز ديمو" → demo URL'],
        ['Each pricing card', 'Primary "تواصل عبر واتساب" + Secondary "احجز ديمو"', 'Reduced to ONE CTA: "احجز ديمو مجاني"'],
        ['Final CTA', 'Two CTAs (start free + WhatsApp)', 'Single primary CTA only'],
        ['Footer social', 'Twitter + WhatsApp', 'Twitter + LinkedIn (WhatsApp removed)'],
    ],
    col_widths=[Cm(3), Cm(7), Cm(7)],
    font_size=9,
)

add_heading(doc, '11.12.4 New Files / Modified Files', level=3)
for cat, items in [
    ('New (CMS layer)', [
        'app/Models/PricingSetting.php',
        'app/Http/Controllers/Landlord/PricingSettingsController.php',
        'resources/views/landlord/pricing-settings.blade.php',
        'database/migrations/2026_07_04_000001_create_pricing_settings_table.php',
        'database/seeders/PricingSettingSeeder.php',
    ]),
    ('Modified', [
        'routes/web.php — added /super-admin/pricing-settings.{index,update}',
        'resources/views/landlord/layout.blade.php — added sidebar item "إعدادات صفحة الأسعار" (with CMS badge)',
        'app/Http/Controllers/PricingController.php — reads via PricingSetting::current() instead of hard-coded',
        'resources/views/pricing/index.blade.php — driven by $settings, WhatsApp removed',
        'resources/views/pricing/layout.blade.php — removed unused .k-btn-whatsapp CSS class',
    ]),
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{cat}:')
    r.bold = True
    for item in items:
        sub = doc.add_paragraph(f'    • {item}', style='List Bullet')

add_heading(doc, '11.12.5 Open the CMS (User Guide)', level=3)
add_para_text = [
    '1. Login to the landlord dashboard: http://localhost/super-admin (or https://superdashboard.yourbrand.com/super-admin).',
    '2. Click "إعدادات صفحة الأسعار" in the left sidebar (it has a small "CMS" badge).',
    '3. By default only the "Hero" section is expanded — click any section title to toggle.',
    '4. Edit any field. Inline help text shows expected format. Hero headline + subtitle accept HTML (so you can use <span class="text-amber-400">word</span> for inline highlights).',
    '5. Hit "حفظ كل التغييرات" at the bottom (sticky save bar).',
    '6. A success flash appears → cache auto-busts → next visitor to /pricing sees updates within seconds.',
    '7. Click "معاينة" to open /pricing in a new tab and confirm the changes.',
]
for s in add_para_text:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '11.12.6 Per-plan Value Props (Customization Hook)', level=3)
add_para_text = [
    'Each plan card displays bullet points (value_props). To customize, two options:',
    '   a. (Quick, no DB change) Edit PricingController::PLAN_VALUE_PROPS_FALLBACK in the controller.',
    '   b. (DB-driven) Add a plan_value_props.{slug} key to the pricing_settings.value JSON (currently read but auto-falls-back).',
    'The PricingController already wires $plan->value_props for both sources.',
]
for s in add_para_text:
    p = doc.add_paragraph(s)
    p.runs[0].font.size = Pt(10)

# Save
doc.save(IN_PATH)
print(f'Updated Part 11 v2 section in DOCX')
print(f'Saved: {IN_PATH}')
print(f'Size: {os.path.getsize(IN_PATH) / 1024:.1f} KB')
